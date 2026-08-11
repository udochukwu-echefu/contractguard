import json
from io import BytesIO, StringIO


CONTRACTGUARD_MODULE_VERSION = 2


def _value(value, fallback="Not identified"):
    return fallback if value in (None, "", []) else str(value)


def build_markdown_report(analysis, source_name, context, notes=""):
    lines = [
        f"# {_value(analysis.get('title') or analysis.get('contract_type'), 'Contract review')}",
        "",
        "> ContractGuard provides education and first-pass triage, not legal advice.",
        "",
        f"- Source: {_value(source_name)}",
        f"- Contract type: {_value(analysis.get('contract_type'))}",
        f"- Contract category: {_value(analysis.get('contract_category'))}",
        f"- Classification confidence: {_value(analysis.get('classification', {}).get('confidence'), 'Needs verification')}",
        f"- Governing law: {_value(analysis.get('governing_law'))}",
        f"- Jurisdiction supplied by reviewer: {'Yes' if analysis.get('jurisdiction_supplied') else 'No'}",
        f"- Reviewing as: {_value(context.get('party_role'))}",
        f"- Review goal: {_value(context.get('goal'))}",
        "",
        "## Executive summary",
        "",
        _value(analysis.get("executive_summary"), "No summary returned."),
        "",
        "## Priority risks",
        "",
    ]
    for index, risk in enumerate(analysis.get("risk_assessment", []), start=1):
        lines.extend(
            [
                f"### {index}. {_value(risk.get('title'), 'Clause finding')} ({_value(risk.get('risk_level'), 'Low')})",
                "",
                _value(risk.get("explanation")),
                "",
                f"- Source clause: {_value(risk.get('clause'))}",
                f"- Consequence: {_value(risk.get('consequence'))}",
                f"- Reviewer action: {_value(risk.get('recommendation'))}",
                f"- Evidence location: {_value(risk.get('citation'))}",
                f"- Applicable category: {_value(risk.get('applicable_category'), 'other')}",
                f"- Confidence: {_value(risk.get('confidence'), 'Low')} (not a measure of legal correctness)",
                f"- Recommendation scope: {_value(risk.get('recommendation_scope'), 'General')}",
                f"- Human review state: {_value(risk.get('human_review_state'), 'No decision')}",
                f"> {_value(risk.get('quote'), 'No exact excerpt returned')}",
                "",
            ]
        )
        if risk.get("suggested_language"):
            lines.extend(["Suggested language to discuss:", "", risk["suggested_language"], ""])

    lines.extend(["## Negotiation plan", ""])
    for item in analysis.get("negotiation_priorities", []):
        lines.extend(
            [
                f"### {_value(item.get('priority'), '•')}. {_value(item.get('title'))}",
                "",
                f"- Ask: {_value(item.get('ask'))}",
                f"- Fallback: {_value(item.get('fallback'))}",
                f"- Evidence: {_value(item.get('citation'))}",
                "",
            ]
        )

    lines.extend(["## Possible protection gaps", ""])
    for item in analysis.get("missing_protections", []):
        lines.extend([f"- **{_value(item.get('issue'))}:** {_value(item.get('explanation'))}"])

    lines.extend(["", "## Obligations", ""])
    for item in analysis.get("obligations", []):
        lines.append(
            f"- **{_value(item.get('party'))}:** {_value(item.get('obligation'))} · {_value(item.get('timing'))} · {_value(item.get('citation'))}"
        )

    if analysis.get("uncertainties"):
        lines.extend(["", "## Items to verify", ""])
        lines.extend(f"- {item}" for item in analysis["uncertainties"])
    if notes.strip():
        lines.extend(["", "## Review notes", "", notes.strip()])
    return "\n".join(lines).strip() + "\n"


def build_docx_report(analysis, source_name, context, notes=""):
    from docx import Document

    document = Document()
    document.add_heading(_value(analysis.get("title") or analysis.get("contract_type"), "Contract review"), 0)
    document.add_paragraph("ContractGuard provides education and first-pass triage, not legal advice.")
    for label, value in [
        ("Source", source_name),
        ("Contract type", analysis.get("contract_type")),
        ("Contract category", analysis.get("contract_category")),
        ("Classification confidence", analysis.get("classification", {}).get("confidence")),
        ("Governing law", analysis.get("governing_law")),
        ("Jurisdiction supplied by reviewer", "Yes" if analysis.get("jurisdiction_supplied") else "No"),
        ("Reviewing as", context.get("party_role")),
        ("Review goal", context.get("goal")),
    ]:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(_value(value))

    document.add_heading("Executive summary", 1)
    document.add_paragraph(_value(analysis.get("executive_summary"), "No summary returned."))
    document.add_heading("Priority risks", 1)
    for risk in analysis.get("risk_assessment", []):
        document.add_heading(f"{_value(risk.get('title'), 'Clause finding')} · {_value(risk.get('risk_level'))}", 2)
        document.add_paragraph(_value(risk.get("explanation")))
        document.add_paragraph(f"Source clause: {_value(risk.get('clause'))}")
        document.add_paragraph(f"Consequence: {_value(risk.get('consequence'))}")
        document.add_paragraph(f"Reviewer action: {_value(risk.get('recommendation'))}")
        document.add_paragraph(f"Evidence: {_value(risk.get('citation'))}: {_value(risk.get('quote'))}")
        document.add_paragraph(f"Category: {_value(risk.get('applicable_category'), 'other')} | Confidence: {_value(risk.get('confidence'), 'Low')} (not legal correctness) | Scope: {_value(risk.get('recommendation_scope'), 'General')} | Human state: {_value(risk.get('human_review_state'), 'No decision')}")
        if risk.get("suggested_language"):
            document.add_paragraph(f"Example language: {risk['suggested_language']}")

    document.add_heading("Negotiation plan", 1)
    for item in analysis.get("negotiation_priorities", []):
        document.add_heading(f"{_value(item.get('priority'), '•')}. {_value(item.get('title'))}", 2)
        document.add_paragraph(f"Ask: {_value(item.get('ask'))}")
        document.add_paragraph(f"Fallback: {_value(item.get('fallback'))}")

    document.add_heading("Obligations", 1)
    for item in analysis.get("obligations", []):
        document.add_paragraph(
            f"{_value(item.get('party'))}: {_value(item.get('obligation'))} · {_value(item.get('timing'))}",
            style="List Bullet",
        )
    if notes.strip():
        document.add_heading("Review notes", 1)
        document.add_paragraph(notes.strip())
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_pdf_report(analysis, source_name, context, notes=""):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    output = BytesIO()
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=16 * mm, bottomMargin=16 * mm)
    story = [Paragraph(_value(analysis.get("title") or analysis.get("contract_type"), "Contract review"), styles["Title"])]
    story.append(Paragraph("ContractGuard provides education and first-pass triage, not legal advice.", styles["Italic"]))
    story.append(Spacer(1, 8))
    story.append(Paragraph(f"Source: {_value(source_name)}<br/>Contract category: {_value(analysis.get('contract_category'))}<br/>Classification confidence: {_value(analysis.get('classification', {}).get('confidence'), 'Needs verification')}<br/>Reviewing as: {_value(context.get('party_role'))}<br/>Governing law: {_value(analysis.get('governing_law'))}<br/>Jurisdiction supplied by reviewer: {'Yes' if analysis.get('jurisdiction_supplied') else 'No'}", styles["BodyText"]))
    story.extend([Spacer(1, 10), Paragraph("Executive summary", styles["Heading1"]), Paragraph(_value(analysis.get("executive_summary")), styles["BodyText"])])
    story.append(Paragraph("Priority risks", styles["Heading1"]))
    for risk in analysis.get("risk_assessment", []):
        story.append(Paragraph(f"{_value(risk.get('title'))} · {_value(risk.get('risk_level'))}", styles["Heading2"]))
        story.append(Paragraph(_value(risk.get("explanation")), styles["BodyText"]))
        story.append(Paragraph(f"Source clause: {_value(risk.get('clause'))}", styles["BodyText"]))
        story.append(Paragraph(f"Consequence: {_value(risk.get('consequence'))}", styles["BodyText"]))
        story.append(Paragraph(f"Reviewer action: {_value(risk.get('recommendation'))}", styles["BodyText"]))
        story.append(Paragraph(f"Evidence: {_value(risk.get('citation'))}: {_value(risk.get('quote'))}", styles["BodyText"]))
        story.append(Paragraph(f"Category: {_value(risk.get('applicable_category'), 'other')} | Confidence: {_value(risk.get('confidence'), 'Low')} (not legal correctness) | Scope: {_value(risk.get('recommendation_scope'), 'General')} | Human state: {_value(risk.get('human_review_state'), 'No decision')}", styles["BodyText"]))
    if notes.strip():
        story.extend([Paragraph("Review notes", styles["Heading1"]), Paragraph(notes.strip(), styles["BodyText"])])
    doc.build(story)
    return output.getvalue()


def build_csv(items):
    import csv

    output = StringIO()
    rows = items or []
    if not rows:
        return ""
    fields = list(dict.fromkeys(key for row in rows for key in row.keys()))
    writer = csv.DictWriter(output, fieldnames=fields)
    writer.writeheader()
    writer.writerows(rows)
    return output.getvalue()


def build_json_report(analysis, context, quality):
    return json.dumps({"review_context": context, "extraction": quality, "analysis": analysis}, indent=2, ensure_ascii=False)
