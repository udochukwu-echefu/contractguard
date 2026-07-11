import json
import os
import re
from functools import lru_cache
from pathlib import Path

from dotenv import load_dotenv

from prompts import COMPARISON_PROMPT, CONTRACT_ANALYSIS_PROMPT, QA_PROMPT, REWRITE_PROMPT


load_dotenv(Path(__file__).with_name(".env"))


def review_context_text(context):
    context = context or {}
    return "\n".join(
        [
            f"User's party: {context.get('party_role') or 'Not specified'}",
            f"Jurisdiction or governing law: {context.get('jurisdiction') or 'Not specified'}",
            f"Review goal: {context.get('goal') or 'Understand before signing'}",
            f"Risk tolerance: {context.get('risk_tolerance') or 'Balanced'}",
        ]
    )


class ContractQAChain:
    """Retrieval QA wrapper that returns both an answer and inspectable evidence."""

    def __init__(self, llm, retriever, review_context=None):
        self.llm = llm
        self.retriever = retriever
        self.review_context = review_context_text(review_context)

    def invoke(self, inputs):
        from langchain_core.output_parsers import StrOutputParser

        user_input = inputs["input"]
        chat_history = inputs.get("chat_history", [])
        retrieval_query = user_input
        if chat_history:
            rewrite_chain = REWRITE_PROMPT | self.llm | StrOutputParser()
            retrieval_query = rewrite_chain.invoke(
                {"input": user_input, "chat_history": chat_history}
            )

        docs = self.retriever.invoke(retrieval_query)
        evidence = []
        context_blocks = []
        for index, doc in enumerate(docs, start=1):
            page = doc.metadata.get("page_label") or doc.metadata.get("page")
            if isinstance(page, int):
                page = page + 1
            location = doc.metadata.get("location") or (f"Page {page}" if page else "Document excerpt")
            excerpt = " ".join(doc.page_content.split())
            evidence.append({"label": f"Source {index}", "location": str(location), "excerpt": excerpt})
            context_blocks.append(f"[Source {index}] {location}\n{doc.page_content}")

        response = (QA_PROMPT | self.llm).invoke(
            {
                "input": user_input,
                "chat_history": chat_history,
                "context": "\n\n".join(context_blocks),
                "review_context": self.review_context,
            }
        )
        return {"answer": response.content, "sources": evidence}


@lru_cache(maxsize=1)
def get_llm():
    from langchain_openai import ChatOpenAI

    api_key = os.environ.get("GROQ_API_KEY")
    if not api_key:
        raise ValueError("GROQ_API_KEY is missing. Add it to your .env file or Streamlit secrets.")
    return ChatOpenAI(
        model="llama-3.3-70b-versatile",
        api_key=api_key,
        base_url="https://api.groq.com/openai/v1",
        temperature=0.1,
        timeout=90,
        max_retries=2,
    )


@lru_cache(maxsize=1)
def get_embeddings():
    """Load the embedding model once per app process, not once per review."""
    from langchain_huggingface import HuggingFaceEmbeddings

    return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")


def _line_numbered(text):
    return "\n".join(f"[L{index}] {line}" for index, line in enumerate(text.splitlines(), start=1))


def _document_quality(full_text, documents, file_type, ocr_used=False):
    characters = len(full_text.strip())
    pages = len(documents) if file_type == "pdf" else 1
    warnings = []
    if characters < 250:
        warnings.append("Very little text was extracted. The document may be scanned or image-based.")
    if "\ufffd" in full_text:
        warnings.append("Some characters could not be decoded correctly.")
    return {
        "characters": characters,
        "pages": pages,
        "file_type": file_type.upper(),
        "ocr_used": ocr_used,
        "quality": "Needs review" if warnings else "Readable",
        "warnings": warnings,
    }


def parse_document(file_path):
    """Return line/page-marked text, citation-ready chunks, and extraction quality."""
    from langchain_core.documents import Document
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        from langchain_community.document_loaders import PyPDFLoader

        documents = PyPDFLoader(file_path).load()
        ocr_used = False
        if len("".join(doc.page_content for doc in documents).strip()) < 250:
            try:
                import pytesseract
                from pdf2image import convert_from_path

                images = convert_from_path(file_path, dpi=200)
                ocr_documents = []
                for index, image in enumerate(images, start=1):
                    ocr_documents.append(
                        Document(
                            page_content=pytesseract.image_to_string(image),
                            metadata={"page_label": index, "location": f"Page {index} (OCR)"},
                        )
                    )
                if len("".join(doc.page_content for doc in ocr_documents).strip()) > 250:
                    documents = ocr_documents
                    ocr_used = True
            except Exception:
                ocr_used = False
        marked_pages = []
        for index, doc in enumerate(documents, start=1):
            doc.metadata["page_label"] = index
            doc.metadata["location"] = f"Page {index}"
            marked_pages.append(f"[PAGE {index}]\n{doc.page_content}")
        full_text = "\n\n".join(marked_pages)
        file_type = "pdf"
    elif suffix == ".docx":
        from docx import Document as DocxDocument

        parsed = DocxDocument(file_path)
        paragraphs = [paragraph.text for paragraph in parsed.paragraphs if paragraph.text.strip()]
        raw_text = "\n".join(paragraphs)
        full_text = _line_numbered(raw_text)
        documents = [Document(page_content=raw_text, metadata={"location": "DOCX paragraphs"})]
        file_type = "docx"
        ocr_used = False
    else:
        raw_text = Path(file_path).read_text(encoding="utf-8", errors="replace")
        full_text = _line_numbered(raw_text)
        documents = [Document(page_content=raw_text, metadata={"location": "Text lines"})]
        file_type = "txt"
        ocr_used = False

    if not full_text.strip():
        return "", [], _document_quality("", documents, file_type, ocr_used)

    splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=140)
    chunks = splitter.split_documents(documents)
    for index, chunk in enumerate(chunks, start=1):
        chunk.metadata["chunk_id"] = index
        if not chunk.metadata.get("location"):
            chunk.metadata["location"] = f"Excerpt {index}"
    return full_text, chunks, _document_quality(full_text, documents, file_type, ocr_used)


def analyze_contract(text, review_context=None):
    llm = get_llm()
    json_llm = llm.bind(response_format={"type": "json_object"})
    response = (CONTRACT_ANALYSIS_PROMPT | json_llm).invoke(
        {
            "contract_text": text[:120_000],
            "review_context": review_context_text(review_context),
        }
    )
    try:
        return json.loads(response.content)
    except json.JSONDecodeError as exc:
        raise ValueError("The model returned an incomplete report. Please retry the analysis.") from exc


def compare_contracts(original, revised, review_context=None):
    llm = get_llm().bind(response_format={"type": "json_object"})
    response = (COMPARISON_PROMPT | llm).invoke(
        {
            "original": original[:60_000],
            "revised": revised[:60_000],
            "review_context": review_context_text(review_context),
        }
    )
    try:
        return json.loads(response.content)
    except json.JSONDecodeError as exc:
        raise ValueError("The model could not create a structured comparison. Please retry.") from exc


def setup_qa_chain(chunks, review_context=None):
    from langchain_chroma import Chroma

    if not chunks:
        raise ValueError("No readable document text was available for Q&A.")
    vectordb = Chroma.from_documents(chunks, get_embeddings())
    retriever = vectordb.as_retriever(search_kwargs={"k": 4})
    return ContractQAChain(get_llm(), retriever, review_context)


def safe_filename(name):
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", name or "contract").strip("-")
    return cleaned or "contract"
