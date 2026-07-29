from __future__ import annotations

import hashlib
import csv
import calendar
import difflib
import io
import json
import math
import os
import re
import tempfile
import secrets
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from uuid import uuid4

from fastapi import HTTPException
from sqlalchemy import func, select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from .config import Settings
from .models import (
    ApprovalRequest,
    Contract,
    ContractComment,
    ContractDecision,
    ContractReview,
    ContractVersion,
    CounterpartyResponse,
    DocumentAsset,
    ExternalShare,
    IntegrationConnection,
    IntegrationImport,
    LifecycleItem,
    Membership,
    NegotiationItem,
    Notification,
    Organization,
    OrganizationInvitation,
    OrganizationSettings,
    PlatformAuditEvent,
    ProcessingJob,
    PublicApiKey,
    SecureIntakeLink,
    User,
    VerificationAssignment,
    VerificationCase,
    VerificationDecision,
    VerificationDocument,
    VerificationReconciliation,
    WebhookDelivery,
    WebhookSubscription,
    WorkflowTask,
    utcnow,
)
from .object_storage import ObjectStore
from .malware import scan_upload
from .schemas import (
    ApprovalRequestResponse,
    ApiKeyCreatedResponse,
    ApiKeyResponse,
    AuditEventResponse,
    ContractActivityResponse,
    ContractCommentResponse,
    ContractDecisionResponse,
    DealPassportResponse,
    ContractResponse,
    ContractVersionResponse,
    CounterpartyResponseResponse,
    ExternalShareResponse,
    IntegrationConnectionResponse,
    IntegrationImportResponse,
    IntegrationProviderResponse,
    IntakeAddressResponse,
    IntakeCreatedResponse,
    InvitationPreviewResponse,
    InvitationResponse,
    JobResponse,
    LifecycleItemResponse,
    MembershipResponse,
    NegotiationItemResponse,
    NegotiationSummaryResponse,
    NotificationResponse,
    OrganizationSettingsResponse,
    PortfolioQuestionResponse,
    PortfolioQuestionSource,
    ReportActivityItem,
    ReportDistributionItem,
    ReportOverviewResponse,
    ReportTimelinePoint,
    ReportWorkloadItem,
    ReviewResponse,
    TaskResponse,
    SecureIntakeLinkResponse,
    VerificationAssignmentResponse,
    VerificationCaseResponse,
    VerificationCaseSummaryResponse,
    VerificationDecisionResponse,
    VerificationDocumentResponse,
    VerificationReconciliationResponse,
    WebhookDeliveryResponse,
    WebhookSubscriptionResponse,
)
from .security import Principal


def json_dump(value: Any) -> str:
    return json.dumps(value if value is not None else {}, ensure_ascii=False, default=str)


def json_load(value: str | None, fallback: Any) -> Any:
    if not value:
        return fallback
    try:
        return json.loads(value)
    except (TypeError, json.JSONDecodeError):
        return fallback


def safe_filename(value: str) -> str:
    name = Path(value or "contract").name
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", name).strip("-.")
    return cleaned or "contract"


VALID_ROLES = {"owner", "admin", "reviewer", "viewer"}
INVITABLE_ROLES = {"admin", "reviewer", "viewer"}
TASK_STATUSES = {"open", "in_progress", "done", "cancelled"}
TASK_PRIORITIES = {"low", "normal", "high"}
TASK_CATEGORIES = {"follow_up", "risk", "obligation", "deadline", "negotiation", "professional_review"}
TASK_SOURCE_KINDS = {"manual", "finding", "obligation", "deadline", "payment", "negotiation"}
VERIFICATION_STATUSES = {
    "pending",
    "in_review",
    "needs_information",
    "approved",
    "escalated",
    "rejected",
    "closed",
}
VERIFICATION_ACTIONS = {"Approve", "Escalate", "Reject"}
VERIFICATION_PRIORITIES = {"low", "normal", "high", "urgent"}
RECONCILIATION_STATUSES = {"matched", "conflict", "needs_review", "resolved"}
VERIFICATION_TRANSITIONS = {
    "pending": {"in_review", "needs_information", "escalated", "closed"},
    "in_review": {"needs_information", "approved", "escalated", "rejected", "closed"},
    "needs_information": {"in_review", "escalated", "closed"},
    "approved": {"in_review", "closed"},
    "escalated": {"in_review", "needs_information", "approved", "rejected", "closed"},
    "rejected": {"in_review", "closed"},
    "closed": {"in_review"},
}
CONTRACT_DECISIONS = {"accept", "change", "escalate", "resolve"}
APPROVAL_STATUSES = {"pending", "approved", "conditionally_approved", "changes_requested", "rejected", "cancelled"}
LIFECYCLE_KINDS = {"renewal", "notice", "obligation", "payment", "post_signature"}
LIFECYCLE_STATUSES = {"active", "completed", "cancelled"}
RECURRENCES = {"none", "weekly", "monthly", "quarterly", "yearly"}
NEGOTIATION_STATUSES = {"proposed", "accepted", "rejected", "unresolved", "resolved"}
NEGOTIATION_CATEGORIES = {"change", "commercial", "legal", "operational", "open_point"}
INTEGRATION_PROVIDERS = {
    "email",
    "google_drive",
    "onedrive",
    "sharepoint",
    "dropbox",
    "slack",
    "telegram",
    "whatsapp",
    "public_api",
}
INTEGRATION_PROVIDER_CATALOG = {
    "email": ("Contract forwarding email", "email", ["document_intake"], "managed"),
    "google_drive": ("Google Drive", "cloud_storage", ["document_import", "folder_watch"], "oauth"),
    "onedrive": ("OneDrive", "cloud_storage", ["document_import", "folder_watch"], "oauth"),
    "sharepoint": ("SharePoint", "cloud_storage", ["document_import", "site_library_watch"], "oauth"),
    "dropbox": ("Dropbox", "cloud_storage", ["document_import", "folder_watch"], "oauth"),
    "slack": ("Slack", "messaging", ["document_intake", "review_notifications"], "oauth"),
    "telegram": ("Telegram", "messaging", ["document_intake", "review_notifications"], "bot"),
    "whatsapp": ("WhatsApp secure links", "messaging", ["secure_links", "review_notifications"], "secure_link"),
    "public_api": ("Public API", "developer", ["document_intake", "status_read", "webhooks"], "api_key"),
}
WEBHOOK_EVENTS = {"contract.created", "contract.review_ready", "contract.review_failed"}
API_KEY_SCOPES = {"contracts:write", "contracts:read", "webhooks:read"}
EMAIL_PATTERN = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


def normalized_role(role: str) -> str:
    return "reviewer" if role == "member" else role


def normalized_email(email: str) -> str:
    value = email.strip().casefold()
    if not EMAIL_PATTERN.fullmatch(value):
        raise HTTPException(status_code=422, detail="Enter a valid email address.")
    return value


def aware(value: datetime) -> datetime:
    return value if value.tzinfo else value.replace(tzinfo=timezone.utc)


def invitation_status(invitation: OrganizationInvitation) -> str:
    if invitation.accepted_at is not None:
        return "accepted"
    if invitation.revoked_at is not None:
        return "revoked"
    if aware(invitation.expires_at) <= utcnow():
        return "expired"
    return "pending"


def email_hint(email: str) -> str:
    local, _, domain = email.partition("@")
    visible = local[:1]
    return f"{visible}{'*' * max(len(local) - 1, 2)}@{domain}"


class PlatformService:
    def __init__(self, session: Session, settings: Settings, object_store: ObjectStore):
        self.session = session
        self.settings = settings
        self.object_store = object_store

    def ensure_user(self, principal: Principal) -> User:
        user = self.session.scalar(select(User).where(User.external_subject == principal.subject))
        if user is None:
            user = User(
                external_subject=principal.subject,
                email=principal.email,
                display_name=principal.display_name,
            )
            self.session.add(user)
            try:
                self.session.commit()
            except IntegrityError:
                self.session.rollback()
                user = self.session.scalar(
                    select(User).where(User.external_subject == principal.subject)
                )
                if user is None:
                    raise
            self.session.refresh(user)
            return user
        changed = False
        if principal.email and user.email != principal.email:
            user.email = principal.email
            changed = True
        if principal.display_name and user.display_name != principal.display_name:
            user.display_name = principal.display_name
            changed = True
        if changed:
            self.session.commit()
        return user

    def create_organization(self, user: User, name: str, slug: str) -> Organization:
        normalized_slug = slug.strip().lower()
        if self.session.scalar(select(Organization.id).where(Organization.slug == normalized_slug)):
            raise HTTPException(status_code=409, detail="That organization slug is already in use.")
        organization = Organization(name=name.strip(), slug=normalized_slug)
        self.session.add(organization)
        self.session.flush()
        self.session.add(Membership(organization_id=organization.id, user_id=user.id, role="owner"))
        self.session.add(OrganizationSettings(organization_id=organization.id))
        self._audit(organization.id, user.id, "organization.created", detail={"name": organization.name})
        try:
            self.session.commit()
        except IntegrityError as exc:
            self.session.rollback()
            raise HTTPException(status_code=409, detail="That organization slug is already in use.") from exc
        self.session.refresh(organization)
        return organization

    def list_organizations(self, user: User) -> list[tuple[Organization, str]]:
        rows = self.session.execute(
            select(Organization, Membership.role)
            .join(Membership, Membership.organization_id == Organization.id)
            .where(Membership.user_id == user.id)
            .order_by(Organization.name.asc())
        ).all()
        return [(row[0], row[1]) for row in rows]

    def organization_settings(self, organization_id: str, user: User) -> OrganizationSettings:
        self.membership(organization_id, user)
        settings = self.session.scalar(
            select(OrganizationSettings).where(OrganizationSettings.organization_id == organization_id)
        )
        if settings is None:
            settings = OrganizationSettings(organization_id=organization_id)
            self.session.add(settings)
            self.session.commit()
            self.session.refresh(settings)
        return settings

    def update_organization_settings(
        self,
        organization_id: str,
        user: User,
        changes: dict[str, Any],
    ) -> OrganizationSettings:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can change workspace settings.",
        )
        settings = self.organization_settings(organization_id, user)
        organization = self.session.get(Organization, organization_id)
        if organization is None:
            raise HTTPException(status_code=404, detail="Workspace not found.")
        if "name" in changes and changes["name"] is not None:
            organization.name = changes["name"].strip()
        for field in (
            "default_retention_days",
            "default_retain_document",
            "default_retain_source_text",
            "notification_review_ready",
            "notification_review_failed",
        ):
            if field in changes and changes[field] is not None:
                setattr(settings, field, changes[field])
        settings.updated_at = utcnow()
        self._audit(
            organization_id,
            user.id,
            "organization.settings_updated",
            detail={"changed_fields": sorted(changes)},
        )
        self.session.commit()
        self.session.refresh(settings)
        return settings

    def organization_settings_response(
        self,
        settings: OrganizationSettings,
    ) -> OrganizationSettingsResponse:
        return OrganizationSettingsResponse(
            organization_id=settings.organization_id,
            name=settings.organization.name,
            slug=settings.organization.slug,
            default_retention_days=settings.default_retention_days,
            default_retain_document=settings.default_retain_document,
            default_retain_source_text=settings.default_retain_source_text,
            notification_review_ready=settings.notification_review_ready,
            notification_review_failed=settings.notification_review_failed,
            updated_at=settings.updated_at,
        )

    def membership(self, organization_id: str, user: User) -> Membership:
        membership = self.session.scalar(
            select(Membership).where(
                Membership.organization_id == organization_id,
                Membership.user_id == user.id,
            )
        )
        if membership is None:
            raise HTTPException(status_code=403, detail="You do not have access to this organization.")
        return membership

    def require_roles(self, organization_id: str, user: User, allowed: set[str], detail: str) -> Membership:
        membership = self.membership(organization_id, user)
        if normalized_role(membership.role) not in allowed:
            raise HTTPException(status_code=403, detail=detail)
        return membership

    def list_members(self, organization_id: str, user: User) -> list[Membership]:
        self.membership(organization_id, user)
        return list(
            self.session.scalars(
                select(Membership)
                .join(User, User.id == Membership.user_id)
                .where(Membership.organization_id == organization_id)
                .order_by(User.display_name.asc(), User.email.asc())
            ).all()
        )

    def list_invitations(self, organization_id: str, user: User) -> list[OrganizationInvitation]:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can view invitations.",
        )
        return list(
            self.session.scalars(
                select(OrganizationInvitation)
                .where(OrganizationInvitation.organization_id == organization_id)
                .order_by(OrganizationInvitation.created_at.desc())
            ).all()
        )

    def create_invitation(
        self,
        organization_id: str,
        user: User,
        email: str,
        role: str,
    ) -> tuple[OrganizationInvitation, str]:
        actor = self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can invite team members.",
        )
        actor_role = normalized_role(actor.role)
        if role not in INVITABLE_ROLES:
            raise HTTPException(status_code=422, detail="Invitations can grant admin, reviewer, or viewer access.")
        if actor_role == "admin" and role == "admin":
            raise HTTPException(status_code=403, detail="Only an owner can invite another administrator.")

        invite_email = normalized_email(email)
        existing_member = self.session.scalar(
            select(Membership)
            .join(User, User.id == Membership.user_id)
            .where(
                Membership.organization_id == organization_id,
                func.lower(User.email) == invite_email,
            )
        )
        if existing_member is not None:
            raise HTTPException(status_code=409, detail="That person is already a member of this organization.")

        existing_invitation = self.session.scalar(
            select(OrganizationInvitation).where(
                OrganizationInvitation.organization_id == organization_id,
                func.lower(OrganizationInvitation.email) == invite_email,
                OrganizationInvitation.accepted_at.is_(None),
                OrganizationInvitation.revoked_at.is_(None),
                OrganizationInvitation.expires_at > utcnow(),
            )
        )
        if existing_invitation is not None:
            raise HTTPException(status_code=409, detail="A pending invitation already exists for that email address.")

        token = secrets.token_urlsafe(32)
        invitation = OrganizationInvitation(
            organization_id=organization_id,
            created_by_user_id=user.id,
            email=invite_email,
            role=role,
            token_hash=hashlib.sha256(token.encode("utf-8")).hexdigest(),
            expires_at=utcnow() + timedelta(days=7),
        )
        self.session.add(invitation)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "invitation.created",
            detail={"invitation_id": invitation.id, "email": invite_email, "role": role},
        )
        self.session.commit()
        self.session.refresh(invitation)
        return invitation, token

    def find_invitation(self, token: str) -> OrganizationInvitation:
        if not token:
            raise HTTPException(status_code=404, detail="Invitation not found.")
        token_hash = hashlib.sha256(token.encode("utf-8")).hexdigest()
        invitation = self.session.scalar(
            select(OrganizationInvitation).where(OrganizationInvitation.token_hash == token_hash)
        )
        if invitation is None:
            raise HTTPException(status_code=404, detail="Invitation not found.")
        return invitation

    def accept_invitation(self, token: str, user: User) -> tuple[Organization, Membership]:
        invitation = self.find_invitation(token)
        status_value = invitation_status(invitation)
        if status_value == "revoked":
            raise HTTPException(status_code=410, detail="This invitation has been revoked.")
        if status_value == "expired":
            raise HTTPException(status_code=410, detail="This invitation has expired.")
        if not user.email or user.email.casefold() != invitation.email.casefold():
            raise HTTPException(
                status_code=403,
                detail="Sign in with the email address that received this invitation.",
            )

        existing = self.session.scalar(
            select(Membership).where(
                Membership.organization_id == invitation.organization_id,
                Membership.user_id == user.id,
            )
        )
        if invitation.accepted_at is not None:
            if existing is None:
                raise HTTPException(status_code=409, detail="This invitation has already been accepted.")
            return invitation.organization, existing

        membership = existing or Membership(
            organization_id=invitation.organization_id,
            user_id=user.id,
            role=invitation.role,
        )
        if existing is None:
            self.session.add(membership)
        invitation.accepted_at = utcnow()
        self._audit(
            invitation.organization_id,
            user.id,
            "invitation.accepted",
            detail={"invitation_id": invitation.id, "role": invitation.role},
        )
        self.session.commit()
        self.session.refresh(membership)
        return invitation.organization, membership

    def revoke_invitation(self, organization_id: str, invitation_id: str, user: User) -> None:
        actor = self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can revoke invitations.",
        )
        invitation = self.session.scalar(
            select(OrganizationInvitation).where(
                OrganizationInvitation.id == invitation_id,
                OrganizationInvitation.organization_id == organization_id,
            )
        )
        if invitation is None:
            raise HTTPException(status_code=404, detail="Invitation not found.")
        if normalized_role(actor.role) == "admin" and invitation.role == "admin":
            raise HTTPException(status_code=403, detail="Only an owner can manage administrator invitations.")
        if invitation.accepted_at is not None:
            raise HTTPException(status_code=409, detail="Accepted invitations cannot be revoked.")
        if invitation.revoked_at is None:
            invitation.revoked_at = utcnow()
            self._audit(
                organization_id,
                user.id,
                "invitation.revoked",
                detail={"invitation_id": invitation.id, "email": invitation.email},
            )
            self.session.commit()

    def update_member_role(
        self,
        organization_id: str,
        membership_id: str,
        role: str,
        user: User,
    ) -> Membership:
        actor = self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can change member roles.",
        )
        if role not in VALID_ROLES:
            raise HTTPException(status_code=422, detail="Unknown organization role.")
        target = self._organization_membership(organization_id, membership_id)
        actor_role = normalized_role(actor.role)
        target_role = normalized_role(target.role)
        if actor_role == "admin" and (target_role in {"owner", "admin"} or role in {"owner", "admin"}):
            raise HTTPException(status_code=403, detail="Only an owner can manage owners or administrators.")
        if target_role == "owner" and role != "owner" and self._owner_count(organization_id) <= 1:
            raise HTTPException(status_code=409, detail="The organization must keep at least one owner.")
        previous_role = target_role
        target.role = role
        self._audit(
            organization_id,
            user.id,
            "membership.role_changed",
            detail={
                "membership_id": target.id,
                "user_id": target.user_id,
                "previous_role": previous_role,
                "role": role,
            },
        )
        self.session.commit()
        self.session.refresh(target)
        return target

    def remove_member(self, organization_id: str, membership_id: str, user: User) -> None:
        actor = self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can remove team members.",
        )
        target = self._organization_membership(organization_id, membership_id)
        actor_role = normalized_role(actor.role)
        target_role = normalized_role(target.role)
        if actor_role == "admin" and target_role in {"owner", "admin"}:
            raise HTTPException(status_code=403, detail="Only an owner can remove owners or administrators.")
        if target_role == "owner" and self._owner_count(organization_id) <= 1:
            raise HTTPException(status_code=409, detail="The organization must keep at least one owner.")
        self._audit(
            organization_id,
            user.id,
            "membership.removed",
            detail={"membership_id": target.id, "user_id": target.user_id, "role": target_role},
        )
        self.session.delete(target)
        self.session.commit()

    def _organization_membership(self, organization_id: str, membership_id: str) -> Membership:
        membership = self.session.scalar(
            select(Membership).where(
                Membership.id == membership_id,
                Membership.organization_id == organization_id,
            )
        )
        if membership is None:
            raise HTTPException(status_code=404, detail="Member not found.")
        return membership

    def _owner_count(self, organization_id: str) -> int:
        return int(
            self.session.scalar(
                select(func.count(Membership.id)).where(
                    Membership.organization_id == organization_id,
                    Membership.role == "owner",
                )
            )
            or 0
        )

    def list_tasks(
        self,
        organization_id: str,
        user: User,
        *,
        status: str | None = None,
        assigned_to_user_id: str | None = None,
        contract_id: str | None = None,
        due_before: datetime | None = None,
        due_after: datetime | None = None,
    ) -> list[WorkflowTask]:
        self.membership(organization_id, user)
        query = select(WorkflowTask).where(WorkflowTask.organization_id == organization_id)
        if status:
            if status not in TASK_STATUSES:
                raise HTTPException(status_code=422, detail="Unknown task status.")
            query = query.where(WorkflowTask.status == status)
        if assigned_to_user_id:
            query = query.where(WorkflowTask.assigned_to_user_id == assigned_to_user_id)
        if contract_id:
            query = query.where(WorkflowTask.contract_id == contract_id)
        if due_before:
            query = query.where(WorkflowTask.due_at <= due_before)
        if due_after:
            query = query.where(WorkflowTask.due_at >= due_after)
        return list(
            self.session.scalars(
                query.order_by(WorkflowTask.due_at.asc().nulls_last(), WorkflowTask.created_at.desc())
            ).all()
        )

    def get_task(self, organization_id: str, task_id: str, user: User) -> WorkflowTask:
        self.membership(organization_id, user)
        task = self.session.scalar(
            select(WorkflowTask).where(
                WorkflowTask.id == task_id,
                WorkflowTask.organization_id == organization_id,
            )
        )
        if task is None:
            raise HTTPException(status_code=404, detail="Task not found.")
        return task

    def create_task(self, organization_id: str, user: User, payload: dict[str, Any]) -> WorkflowTask:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers have read-only access and cannot create tasks.",
        )
        contract = self._task_contract(organization_id, payload.get("contract_id"))
        assignee = self._task_assignee(organization_id, payload.get("assigned_to_user_id"))
        source_kind = payload.get("source_kind", "manual")
        if source_kind not in TASK_SOURCE_KINDS:
            raise HTTPException(status_code=422, detail="Unknown task source.")
        task = WorkflowTask(
            organization_id=organization_id,
            contract_id=contract.id if contract else None,
            created_by_user_id=user.id,
            assigned_to_user_id=assignee.id if assignee else None,
            title=payload["title"].strip(),
            description=payload.get("description", "").strip(),
            category=payload.get("category", "follow_up"),
            priority=payload.get("priority", "normal"),
            status=payload.get("status", "open"),
            due_at=payload.get("due_at"),
            source_kind=source_kind,
            source_reference_json=json_dump(payload.get("source_reference", {})),
            completed_at=utcnow() if payload.get("status") == "done" else None,
        )
        self.session.add(task)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "task.created",
            task.contract_id,
            {"task_id": task.id, "title": task.title, "assigned_to_user_id": task.assigned_to_user_id},
        )
        self.session.commit()
        self.session.refresh(task)
        return task

    def update_task(
        self,
        organization_id: str,
        task_id: str,
        user: User,
        changes: dict[str, Any],
    ) -> WorkflowTask:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers have read-only access and cannot change tasks.",
        )
        task = self.get_task(organization_id, task_id, user)
        if "contract_id" in changes:
            contract = self._task_contract(organization_id, changes["contract_id"])
            task.contract_id = contract.id if contract else None
        if "assigned_to_user_id" in changes:
            assignee = self._task_assignee(organization_id, changes["assigned_to_user_id"])
            task.assigned_to_user_id = assignee.id if assignee else None
        for field in ("title", "description", "category", "priority", "due_at"):
            if field in changes:
                value = changes[field]
                if value is None and field != "due_at":
                    continue
                if field in {"title", "description"} and isinstance(value, str):
                    value = value.strip()
                setattr(task, field, value)
        if "status" in changes and changes["status"] is not None:
            task.status = changes["status"]
            task.completed_at = utcnow() if task.status == "done" else None
        self._audit(
            organization_id,
            user.id,
            "task.updated",
            task.contract_id,
            {"task_id": task.id, "changed_fields": sorted(changes)},
        )
        self.session.commit()
        self.session.refresh(task)
        return task

    def delete_task(self, organization_id: str, task_id: str, user: User) -> None:
        membership = self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers have read-only access and cannot delete tasks.",
        )
        task = self.get_task(organization_id, task_id, user)
        if normalized_role(membership.role) == "reviewer" and task.created_by_user_id != user.id:
            raise HTTPException(status_code=403, detail="Reviewers can delete only tasks they created.")
        self._audit(
            organization_id,
            user.id,
            "task.deleted",
            task.contract_id,
            {"task_id": task.id, "title": task.title},
        )
        self.session.delete(task)
        self.session.commit()

    def _task_contract(self, organization_id: str, contract_id: str | None) -> Contract | None:
        if not contract_id:
            return None
        contract = self.session.scalar(
            select(Contract).where(
                Contract.id == contract_id,
                Contract.organization_id == organization_id,
            )
        )
        if contract is None:
            raise HTTPException(status_code=422, detail="Choose a contract from this workspace.")
        return contract

    def _task_assignee(self, organization_id: str, user_id: str | None) -> User | None:
        if not user_id:
            return None
        membership = self.session.scalar(
            select(Membership).where(
                Membership.organization_id == organization_id,
                Membership.user_id == user_id,
            )
        )
        if membership is None:
            raise HTTPException(status_code=422, detail="Choose a member of this workspace.")
        return membership.user

    def list_verification_cases(
        self,
        organization_id: str,
        user: User,
        *,
        status: str | None = None,
        suggested_action: str | None = None,
        priority: str | None = None,
        assigned_to_user_id: str | None = None,
        search: str | None = None,
    ) -> list[VerificationCase]:
        self.membership(organization_id, user)
        query = select(VerificationCase).where(VerificationCase.organization_id == organization_id)
        if status:
            if status not in VERIFICATION_STATUSES:
                raise HTTPException(status_code=422, detail="Unknown verification status.")
            query = query.where(VerificationCase.status == status)
        if suggested_action:
            if suggested_action not in VERIFICATION_ACTIONS:
                raise HTTPException(status_code=422, detail="Unknown verification recommendation.")
            query = query.where(VerificationCase.suggested_action == suggested_action)
        if priority:
            if priority not in VERIFICATION_PRIORITIES:
                raise HTTPException(status_code=422, detail="Unknown verification priority.")
            query = query.where(VerificationCase.priority == priority)
        if assigned_to_user_id:
            query = query.where(VerificationCase.assigned_to_user_id == assigned_to_user_id)
        if search and search.strip():
            term = f"%{search.strip().lower()}%"
            query = query.where(
                func.lower(VerificationCase.applicant_name).like(term)
                | func.lower(VerificationCase.reference).like(term)
                | func.lower(VerificationCase.applicant_email).like(term)
            )
        return list(
            self.session.scalars(
                query.order_by(
                    (VerificationCase.status != "pending").asc(),
                    (VerificationCase.priority == "urgent").desc(),
                    (VerificationCase.priority == "high").desc(),
                    VerificationCase.risk_score.desc(),
                    VerificationCase.submitted_at.asc(),
                )
            ).all()
        )

    def _verification_document_uploads(
        self,
        case: VerificationCase,
        uploaded_by_user_id: str | None,
        uploads: list[dict[str, Any]],
    ) -> list[VerificationDocument]:
        documents: list[VerificationDocument] = []
        stored_keys: list[str] = []
        try:
            for upload in uploads:
                filename = safe_filename(str(upload.get("original_name") or "onboarding-document"))
                suffix = Path(filename).suffix.lower()
                if suffix not in self.settings.allowed_extension_set:
                    raise HTTPException(status_code=415, detail="Supported file types are PDF, DOCX, and TXT.")
                data = upload.get("data") or b""
                if not data:
                    raise HTTPException(status_code=400, detail=f"{filename} is empty.")
                if len(data) > self.settings.max_upload_bytes:
                    raise HTTPException(status_code=413, detail=f"{filename} exceeds the 25 MB limit.")
                scan_upload(self.settings, filename, data)
                document_id = str(uuid4())
                storage_key = f"{case.organization_id}/verification/{case.id}/{document_id}{suffix}"
                content_type = str(upload.get("content_type") or "application/octet-stream")
                self.object_store.put(storage_key, data, content_type)
                stored_keys.append(storage_key)
                document = VerificationDocument(
                    id=document_id,
                    organization_id=case.organization_id,
                    verification_case_id=case.id,
                    uploaded_by_user_id=uploaded_by_user_id,
                    document_type=str(upload.get("document_type") or "supporting_document")[:64],
                    original_name=filename,
                    storage_key=storage_key,
                    content_type=content_type[:255],
                    size_bytes=len(data),
                    sha256=hashlib.sha256(data).hexdigest(),
                    status="received",
                    scan_status="clean",
                    extraction_status="pending",
                    expires_at=case.expires_at,
                )
                self.session.add(document)
                documents.append(document)
            self.session.flush()
            return documents
        except Exception:
            for storage_key in stored_keys:
                self.object_store.delete(storage_key)
            raise

    def create_verification_case(
        self,
        organization_id: str,
        user: User,
        *,
        applicant_name: str,
        applicant_email: str,
        reference: str,
        priority: str,
        assigned_to_user_id: str | None,
        due_at: datetime | None,
        retention_days: int,
        intake_channel: str,
        uploads: list[dict[str, Any]],
    ) -> VerificationCase:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers have read-only access and cannot create verification cases.",
        )
        name = applicant_name.strip()
        if len(name) < 2:
            raise HTTPException(status_code=422, detail="Enter the applicant name.")
        if applicant_email.strip():
            normalized_email(applicant_email)
        if priority not in VERIFICATION_PRIORITIES:
            raise HTTPException(status_code=422, detail="Choose a valid priority.")
        if retention_days not in {7, 30, 90, 365}:
            raise HTTPException(status_code=422, detail="Retention must be 7, 30, 90, or 365 days.")
        assignee = self._task_assignee(organization_id, assigned_to_user_id)
        case_reference = (reference.strip() or f"VC-{utcnow():%Y%m%d}-{uuid4().hex[:8].upper()}")[:128]
        if self.session.scalar(
            select(VerificationCase.id).where(
                VerificationCase.organization_id == organization_id,
                VerificationCase.reference == case_reference,
            )
        ):
            raise HTTPException(status_code=409, detail="A verification case already uses this reference.")
        now = utcnow()
        case = VerificationCase(
            organization_id=organization_id,
            seeded_by_user_id=user.id,
            reference=case_reference,
            applicant_name=name[:255],
            applicant_email=applicant_email.strip().lower()[:320],
            status="pending",
            priority=priority,
            assigned_to_user_id=assignee.id if assignee else None,
            intake_channel=intake_channel[:64],
            risk_score=0,
            suggested_action="Escalate",
            finding_count=0,
            document_count=0,
            average_confidence=0,
            submitted_at=now,
            source_json=json_dump({"application": {"name": name, "email": applicant_email.strip().lower()}, "documents": []}),
            evaluation_json=json_dump(
                {
                    "summary": "Documents received. Evidence reconciliation is awaiting reviewer input.",
                    "reasoning": "No automated verification decision has been made.",
                    "findings": [],
                    "field_matrix": [],
                    "generated_at": now.isoformat(),
                }
            ),
            synthetic=False,
            retention_days=retention_days,
            due_at=due_at,
            expires_at=now + timedelta(days=retention_days),
        )
        self.session.add(case)
        self.session.flush()
        documents = self._verification_document_uploads(case, user.id, uploads)
        case.document_count = len(documents)
        if assignee:
            self.session.add(
                VerificationAssignment(
                    organization_id=organization_id,
                    verification_case_id=case.id,
                    assigned_to_user_id=assignee.id,
                    assigned_by_user_id=user.id,
                    note="Assigned during case creation.",
                )
            )
        self._audit(
            organization_id,
            user.id,
            "verification.case_created",
            detail={"reference": case.reference, "document_count": len(documents), "intake_channel": intake_channel},
            verification_case_id=case.id,
        )
        self.session.commit()
        self.session.refresh(case)
        return case

    def update_verification_case(
        self,
        organization_id: str,
        case_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> VerificationCase:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot update verification cases.",
        )
        case = self.get_verification_case(organization_id, case_id, user)
        changed: dict[str, Any] = {}
        if payload.get("priority") is not None:
            priority = payload["priority"]
            if priority not in VERIFICATION_PRIORITIES:
                raise HTTPException(status_code=422, detail="Choose a valid priority.")
            case.priority = priority
            changed["priority"] = priority
        if "due_at" in payload:
            case.due_at = payload["due_at"]
            changed["due_at"] = payload["due_at"]
        if payload.get("status") is not None and payload["status"] != case.status:
            next_status = payload["status"]
            if next_status not in VERIFICATION_TRANSITIONS.get(case.status, set()):
                raise HTTPException(
                    status_code=409,
                    detail=f"A case cannot move from {case.status} to {next_status}.",
                )
            changed["status"] = {"from": case.status, "to": next_status}
            case.status = next_status
            case.closed_at = utcnow() if next_status == "closed" else None
        if not changed:
            return case
        self._audit(
            organization_id,
            user.id,
            "verification.case_updated",
            detail={"reference": case.reference, "changes": changed},
            verification_case_id=case.id,
        )
        self.session.commit()
        self.session.refresh(case)
        return case

    def assign_verification_case(
        self,
        organization_id: str,
        case_id: str,
        user: User,
        *,
        assigned_to_user_id: str | None,
        note: str,
    ) -> VerificationAssignment:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot assign verification cases.",
        )
        case = self.get_verification_case(organization_id, case_id, user)
        assignee = self._task_assignee(organization_id, assigned_to_user_id)
        assignment = VerificationAssignment(
            organization_id=organization_id,
            verification_case_id=case.id,
            assigned_to_user_id=assignee.id if assignee else None,
            assigned_by_user_id=user.id,
            note=note.strip()[:2000],
        )
        case.assigned_to_user_id = assignee.id if assignee else None
        if assignee and case.status == "pending":
            case.status = "in_review"
        self.session.add(assignment)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "verification.case_assigned" if assignee else "verification.case_unassigned",
            detail={
                "reference": case.reference,
                "assigned_to_user_id": assignee.id if assignee else None,
                "note": assignment.note,
            },
            verification_case_id=case.id,
        )
        self.session.commit()
        self.session.refresh(assignment)
        return assignment

    def upsert_verification_reconciliation(
        self,
        organization_id: str,
        case_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> VerificationReconciliation:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot reconcile verification evidence.",
        )
        case = self.get_verification_case(organization_id, case_id, user)
        field_name = payload["field_name"].strip().lower().replace(" ", "_")
        status = payload["status"]
        if status not in RECONCILIATION_STATUSES:
            raise HTTPException(status_code=422, detail="Choose a valid reconciliation status.")
        record = self.session.scalar(
            select(VerificationReconciliation).where(
                VerificationReconciliation.verification_case_id == case.id,
                VerificationReconciliation.field_name == field_name,
            )
        )
        if record is None:
            record = VerificationReconciliation(
                organization_id=organization_id,
                verification_case_id=case.id,
                field_name=field_name,
            )
            self.session.add(record)
        record.canonical_value = payload.get("canonical_value", "").strip()
        record.status = status
        record.sources_json = json_dump(payload.get("sources", []))
        record.resolution_note = payload.get("resolution_note", "").strip()
        if status in {"matched", "resolved"}:
            record.resolved_by_user_id = user.id
            record.resolved_at = utcnow()
        else:
            record.resolved_by_user_id = None
            record.resolved_at = None
        self.session.flush()
        case.finding_count = self.session.scalar(
            select(func.count(VerificationReconciliation.id)).where(
                VerificationReconciliation.verification_case_id == case.id,
                VerificationReconciliation.status.in_({"conflict", "needs_review"}),
            )
        ) or 0
        self._audit(
            organization_id,
            user.id,
            "verification.evidence_reconciled",
            detail={"field_name": field_name, "status": status, "reference": case.reference},
            verification_case_id=case.id,
        )
        self.session.commit()
        self.session.refresh(record)
        return record

    def review_verification_document(
        self,
        organization_id: str,
        case_id: str,
        document_id: str,
        user: User,
        *,
        scan_status: str,
        extraction_status: str,
        extracted_fields: dict[str, Any],
        confidence: int,
    ) -> VerificationDocument:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot update onboarding documents.",
        )
        case = self.get_verification_case(organization_id, case_id, user)
        document = self.session.scalar(
            select(VerificationDocument).where(
                VerificationDocument.id == document_id,
                VerificationDocument.verification_case_id == case.id,
            )
        )
        if document is None:
            raise HTTPException(status_code=404, detail="Verification document not found.")
        if scan_status not in {"pending", "clean", "rejected"}:
            raise HTTPException(status_code=422, detail="Choose pending, clean, or rejected scan status.")
        if extraction_status not in {"pending", "processing", "ready", "failed"}:
            raise HTTPException(status_code=422, detail="Choose a valid extraction status.")
        document.scan_status = scan_status
        document.extraction_status = extraction_status
        document.extracted_fields_json = json_dump(extracted_fields)
        document.confidence = min(max(confidence, 0), 100)
        document.status = "rejected" if scan_status == "rejected" else "ready" if scan_status == "clean" and extraction_status == "ready" else "received"
        ready_confidences = [
            item.confidence
            for item in case.documents
            if item.id == document.id or item.status == "ready"
        ]
        if document.status == "ready" and document.confidence not in ready_confidences:
            ready_confidences.append(document.confidence)
        case.average_confidence = round(sum(ready_confidences) / len(ready_confidences)) if ready_confidences else 0
        self._audit(
            organization_id,
            user.id,
            "verification.document_reviewed",
            detail={
                "document_id": document.id,
                "scan_status": scan_status,
                "extraction_status": extraction_status,
                "confidence": document.confidence,
            },
            verification_case_id=case.id,
        )
        self.session.commit()
        self.session.refresh(document)
        return document

    @staticmethod
    def secure_intake_status(link: SecureIntakeLink) -> str:
        if link.revoked_at is not None:
            return "revoked"
        if aware(link.expires_at) <= utcnow():
            return "expired"
        if link.upload_count >= link.max_uploads:
            return "complete"
        return "active"

    def create_secure_intake_link(
        self,
        organization_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> tuple[SecureIntakeLink, str]:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot create secure intake links.",
        )
        token = f"ll_intake_{secrets.token_urlsafe(32)}"
        link = SecureIntakeLink(
            organization_id=organization_id,
            created_by_user_id=user.id,
            token_hash=hashlib.sha256(token.encode("utf-8")).hexdigest(),
            token_prefix=token[:14],
            channel=payload["channel"],
            recipient_name=payload.get("recipient_name", "").strip(),
            recipient_email=payload.get("recipient_email", "").strip().lower(),
            recipient_phone_hint=payload.get("recipient_phone_hint", "").strip(),
            applicant_name=payload["applicant_name"].strip(),
            message=payload.get("message", "").strip(),
            max_uploads=payload["max_uploads"],
            retention_days=payload["retention_days"],
            expires_at=utcnow() + timedelta(days=payload["expires_in_days"]),
        )
        self.session.add(link)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "verification.intake_link_created",
            detail={"intake_link_id": link.id, "channel": link.channel, "expires_at": link.expires_at},
        )
        self.session.commit()
        self.session.refresh(link)
        return link, token

    def list_secure_intake_links(self, organization_id: str, user: User) -> list[SecureIntakeLink]:
        self.membership(organization_id, user)
        return list(
            self.session.scalars(
                select(SecureIntakeLink)
                .where(SecureIntakeLink.organization_id == organization_id)
                .order_by(SecureIntakeLink.created_at.desc())
            ).all()
        )

    def revoke_secure_intake_link(
        self,
        organization_id: str,
        link_id: str,
        user: User,
    ) -> None:
        self.require_roles(organization_id, user, {"owner", "admin", "reviewer"}, "Viewers cannot revoke intake links.")
        link = self.session.scalar(
            select(SecureIntakeLink).where(
                SecureIntakeLink.id == link_id,
                SecureIntakeLink.organization_id == organization_id,
            )
        )
        if link is None:
            raise HTTPException(status_code=404, detail="Secure intake link not found.")
        link.revoked_at = utcnow()
        self._audit(
            organization_id,
            user.id,
            "verification.intake_link_revoked",
            detail={"intake_link_id": link.id},
            verification_case_id=link.verification_case_id,
        )
        self.session.commit()

    def resolve_secure_intake_link(self, token: str, *, require_active: bool = True) -> SecureIntakeLink:
        link = self.session.scalar(
            select(SecureIntakeLink).where(
                SecureIntakeLink.token_hash == hashlib.sha256(token.encode("utf-8")).hexdigest()
            )
        )
        if link is None:
            raise HTTPException(status_code=404, detail="Secure intake link not found.")
        if require_active and self.secure_intake_status(link) != "active":
            raise HTTPException(status_code=410, detail="This secure intake link is no longer active.")
        return link

    def upload_secure_intake_documents(
        self,
        token: str,
        uploads: list[dict[str, Any]],
    ) -> tuple[VerificationCase, list[VerificationDocument]]:
        link = self.resolve_secure_intake_link(token)
        if not uploads:
            raise HTTPException(status_code=400, detail="Attach at least one onboarding document.")
        if link.upload_count + len(uploads) > link.max_uploads:
            raise HTTPException(status_code=409, detail="This upload exceeds the link's remaining document allowance.")
        case = self.session.get(VerificationCase, link.verification_case_id) if link.verification_case_id else None
        if case is None:
            organization = self.session.get(Organization, link.organization_id)
            case = VerificationCase(
                organization_id=link.organization_id,
                seeded_by_user_id=link.created_by_user_id,
                reference=f"VC-{utcnow():%Y%m%d}-{uuid4().hex[:8].upper()}",
                applicant_name=link.applicant_name,
                applicant_email=link.recipient_email,
                status="pending",
                priority="normal",
                intake_channel=link.channel,
                risk_score=0,
                suggested_action="Escalate",
                finding_count=0,
                document_count=0,
                average_confidence=0,
                submitted_at=utcnow(),
                source_json=json_dump(
                    {"application": {"name": link.applicant_name, "email": link.recipient_email}, "documents": []}
                ),
                evaluation_json=json_dump(
                    {
                        "summary": "Secure onboarding documents received.",
                        "reasoning": "Evidence reconciliation and a human decision are required.",
                        "findings": [],
                        "field_matrix": [],
                        "generated_at": utcnow().isoformat(),
                        "organization": organization.name if organization else "",
                    }
                ),
                synthetic=False,
                retention_days=link.retention_days,
                expires_at=utcnow() + timedelta(days=link.retention_days),
            )
            self.session.add(case)
            self.session.flush()
            link.verification_case_id = case.id
        documents = self._verification_document_uploads(case, None, uploads)
        link.upload_count += len(documents)
        link.last_used_at = utcnow()
        case.document_count = len(case.documents)
        self._audit(
            link.organization_id,
            None,
            "verification.secure_documents_received",
            detail={"intake_link_id": link.id, "document_ids": [item.id for item in documents]},
            verification_case_id=case.id,
        )
        self.session.commit()
        self.session.refresh(case)
        return case, documents

    def get_verification_case(
        self,
        organization_id: str,
        case_id: str,
        user: User,
    ) -> VerificationCase:
        self.membership(organization_id, user)
        case = self.session.scalar(
            select(VerificationCase).where(
                VerificationCase.id == case_id,
                VerificationCase.organization_id == organization_id,
            )
        )
        if case is None:
            raise HTTPException(status_code=404, detail="Verification case not found.")
        return case

    def list_verification_audit_events(
        self,
        organization_id: str,
        case_id: str,
        user: User,
    ) -> list[AuditEventResponse]:
        case = self.get_verification_case(organization_id, case_id, user)
        rows = self.session.execute(
            select(PlatformAuditEvent, User.display_name, User.email)
            .outerjoin(User, User.id == PlatformAuditEvent.actor_user_id)
            .where(
                PlatformAuditEvent.organization_id == organization_id,
                PlatformAuditEvent.verification_case_id == case.id,
            )
            .order_by(PlatformAuditEvent.created_at.desc())
        ).all()
        return [
            AuditEventResponse(
                id=event.id,
                action=event.action,
                detail=json_load(event.detail_json, {}),
                actor_user_id=event.actor_user_id,
                actor_name=actor_name or "Secure intake",
                actor_email=actor_email or "",
                contract_id=event.contract_id,
                verification_case_id=event.verification_case_id,
                created_at=event.created_at,
            )
            for event, actor_name, actor_email in rows
        ]

    def bootstrap_verification_cases(
        self,
        organization_id: str,
        user: User,
    ) -> list[VerificationCase]:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers have read-only access and cannot add demonstration cases.",
        )
        from kyc import evaluate_case, get_cases

        existing_references = set(
            self.session.scalars(
                select(VerificationCase.reference).where(
                    VerificationCase.organization_id == organization_id
                )
            ).all()
        )
        created: list[VerificationCase] = []
        for source in get_cases():
            if source["id"] in existing_references:
                continue
            evaluation_date = date.fromisoformat(source["submitted_at"][:10])
            evaluation = evaluate_case(source, today=evaluation_date)
            submitted_at = datetime.fromisoformat(source["submitted_at"])
            if submitted_at.tzinfo is None:
                submitted_at = submitted_at.replace(tzinfo=timezone.utc)
            case = VerificationCase(
                organization_id=organization_id,
                seeded_by_user_id=user.id,
                reference=source["id"],
                applicant_name=source["applicant"],
                status="pending",
                risk_score=evaluation["score"],
                suggested_action=evaluation["suggested_action"],
                finding_count=len(evaluation["findings"]),
                document_count=evaluation["document_count"],
                average_confidence=round(evaluation["average_confidence"] * 100),
                submitted_at=submitted_at,
                source_json=json_dump(source),
                evaluation_json=json_dump(evaluation),
                synthetic=True,
            )
            self.session.add(case)
            created.append(case)
        if created:
            self.session.flush()
            self._audit(
                organization_id,
                user.id,
                "verification.cases_bootstrapped",
                detail={"case_ids": [item.id for item in created], "count": len(created), "synthetic": True},
            )
            self.session.commit()
        return self.list_verification_cases(organization_id, user)

    def record_verification_decision(
        self,
        organization_id: str,
        case_id: str,
        user: User,
        *,
        decision: str,
        rationale: str,
    ) -> VerificationDecision:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers have read-only access and cannot record verification decisions.",
        )
        case = self.get_verification_case(organization_id, case_id, user)
        membership = self.membership(organization_id, user)
        if (
            case.assigned_to_user_id
            and case.assigned_to_user_id != user.id
            and normalized_role(membership.role) not in {"owner", "admin"}
        ):
            raise HTTPException(status_code=403, detail="This case is assigned to another reviewer.")
        cleaned_rationale = rationale.strip()
        if len(cleaned_rationale) < 10:
            raise HTTPException(status_code=422, detail="Explain the evidence supporting this decision.")
        if decision not in VERIFICATION_ACTIONS:
            raise HTTPException(status_code=422, detail="Choose approve, escalate, or reject.")
        if decision == "Approve" and any(
            item.status in {"conflict", "needs_review"} for item in case.reconciliations
        ):
            raise HTTPException(status_code=409, detail="Resolve every evidence conflict before approving this case.")
        event = VerificationDecision(
            organization_id=organization_id,
            verification_case_id=case.id,
            reviewer_user_id=user.id,
            decision=decision,
            rationale=cleaned_rationale,
            recommended_action=case.suggested_action,
        )
        case.status = {"Approve": "approved", "Escalate": "escalated", "Reject": "rejected"}[decision]
        case.closed_at = None
        self.session.add(event)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "verification.decision_recorded",
            detail={
                "verification_case_id": case.id,
                "reference": case.reference,
                "decision_id": event.id,
                "decision": decision,
                "recommended_action": case.suggested_action,
                "overrode_recommendation": decision != case.suggested_action,
            },
            verification_case_id=case.id,
        )
        self.session.commit()
        self.session.refresh(event)
        return event

    def create_contract(
        self,
        organization_id: str,
        user: User,
        *,
        original_name: str,
        content_type: str,
        data: bytes,
        title: str,
        counterparty: str,
        contract_type: str,
        review_context: dict[str, Any],
        retain_document: bool,
        retain_source_text: bool,
        retention_days: int,
    ) -> tuple[Contract, DocumentAsset, ProcessingJob]:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers have read-only access and cannot upload contracts.",
        )
        filename = safe_filename(original_name)
        suffix = Path(filename).suffix.lower()
        if suffix not in self.settings.allowed_extension_set:
            raise HTTPException(status_code=415, detail="Supported file types are PDF, DOCX, and TXT.")
        if not data:
            raise HTTPException(status_code=400, detail="The uploaded document is empty.")
        if len(data) > self.settings.max_upload_bytes:
            raise HTTPException(status_code=413, detail="The uploaded document exceeds the 25 MB limit.")
        scan_upload(self.settings, filename, data)
        if retention_days not in {7, 30, 90, 365}:
            raise HTTPException(status_code=422, detail="Retention must be 7, 30, 90, or 365 days.")

        now = utcnow()
        contract = Contract(
            organization_id=organization_id,
            created_by_user_id=user.id,
            title=(title or Path(filename).stem).strip()[:512],
            source_name=filename,
            counterparty=counterparty.strip()[:255],
            contract_type=(contract_type or "Unknown").strip()[:255],
            status="processing",
            review_context_json=json_dump(review_context),
            retain_document=retain_document,
            retain_source_text=retain_source_text,
            retention_days=retention_days,
            expires_at=now + timedelta(days=retention_days),
        )
        self.session.add(contract)
        self.session.flush()
        asset_id = str(uuid4())
        storage_key = f"{organization_id}/{contract.id}/{asset_id}{suffix}"
        asset = DocumentAsset(
            id=asset_id,
            organization_id=organization_id,
            contract_id=contract.id,
            storage_key=storage_key,
            original_name=filename,
            content_type=content_type or "application/octet-stream",
            size_bytes=len(data),
            sha256=hashlib.sha256(data).hexdigest(),
        )
        version = ContractVersion(
            organization_id=organization_id,
            contract_id=contract.id,
            document_asset_id=asset.id,
            uploaded_by_user_id=user.id,
            version_number=1,
            label="Original",
            notes="Initial uploaded document.",
            source_name=filename,
            sha256=asset.sha256,
            size_bytes=len(data),
            comparison_json=json_dump(
                {
                    "compared_to_version_id": None,
                    "added": [],
                    "removed": [],
                    "changed_summary": "Original uploaded document.",
                    "added_count": 0,
                    "removed_count": 0,
                }
            ),
        )
        job = ProcessingJob(
            organization_id=organization_id,
            contract_id=contract.id,
            document_asset_id=asset.id,
            status="queued",
            progress_step="Waiting for a review worker",
        )
        self.session.add_all((asset, version, job))
        self._audit(
            organization_id,
            user.id,
            "contract.created",
            contract.id,
            {"source_name": filename, "size_bytes": len(data)},
        )
        self._enqueue_webhooks(
            organization_id,
            "contract.created",
            contract.id,
            {
                "contract_id": contract.id,
                "title": contract.title,
                "source_name": filename,
                "status": contract.status,
            },
        )
        try:
            self.object_store.put(storage_key, data, asset.content_type)
            self.session.commit()
        except Exception:
            self.session.rollback()
            self.object_store.delete(storage_key)
            raise
        self.session.refresh(contract)
        self.session.refresh(asset)
        self.session.refresh(job)
        return contract, asset, job

    def list_integration_connections(
        self,
        organization_id: str,
        user: User,
        provider: str | None = None,
    ) -> list[IntegrationConnection]:
        self.membership(organization_id, user)
        query = select(IntegrationConnection).where(IntegrationConnection.organization_id == organization_id)
        if provider:
            if provider not in INTEGRATION_PROVIDERS:
                raise HTTPException(status_code=422, detail="Unknown integration provider.")
            query = query.where(IntegrationConnection.provider == provider)
        return list(
            self.session.scalars(query.order_by(IntegrationConnection.created_at.desc())).all()
        )

    def integration_providers(self, organization_id: str, user: User) -> list[IntegrationProviderResponse]:
        self.membership(organization_id, user)
        configured = set(
            self.session.scalars(
                select(IntegrationConnection.provider).where(
                    IntegrationConnection.organization_id == organization_id,
                    IntegrationConnection.status == "active",
                )
            ).all()
        )
        return [
            IntegrationProviderResponse(
                provider=provider,
                display_name=details[0],
                category=details[1],
                capabilities=details[2],
                connection_mode=details[3],
                configured=provider in configured or provider in {"email", "public_api", "whatsapp"},
            )
            for provider, details in INTEGRATION_PROVIDER_CATALOG.items()
        ]

    def intake_email_address(self, organization_id: str, user: User) -> IntakeAddressResponse:
        organization = self.membership(organization_id, user).organization
        domain = self.settings.intake_email_domain.strip().lower()
        return IntakeAddressResponse(
            address=f"contracts+{organization.slug}@{domain}",
            enabled=bool(domain),
            instructions="Forward one PDF, DOCX, or TXT contract per message. The original sender and message ID are retained in the intake audit record.",
        )

    def create_integration_connection(
        self,
        organization_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> IntegrationConnection:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can manage integrations.",
        )
        provider = payload["provider"]
        if provider not in INTEGRATION_PROVIDERS:
            raise HTTPException(status_code=422, detail="Unknown integration provider.")
        settings = payload.get("settings", {})
        forbidden = {
            key
            for key in settings
            if any(fragment in key.lower() for fragment in ("secret", "token", "password", "private_key"))
        }
        if forbidden:
            raise HTTPException(
                status_code=422,
                detail="Store connector credentials in deployment secrets, not integration settings.",
            )
        provider_defaults = INTEGRATION_PROVIDER_CATALOG[provider]
        capabilities = payload.get("capabilities") or provider_defaults[2]
        connection = IntegrationConnection(
            organization_id=organization_id,
            created_by_user_id=user.id,
            provider=provider,
            display_name=payload["display_name"].strip(),
            external_account_id=payload.get("external_account_id", "").strip(),
            status="active",
            capabilities_json=json_dump(capabilities),
            settings_json=json_dump(settings),
        )
        self.session.add(connection)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "integration.connected",
            detail={"connection_id": connection.id, "provider": provider},
        )
        self.session.commit()
        self.session.refresh(connection)
        return connection

    def revoke_integration_connection(
        self,
        organization_id: str,
        connection_id: str,
        user: User,
    ) -> IntegrationConnection:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can manage integrations.",
        )
        connection = self.session.scalar(
            select(IntegrationConnection).where(
                IntegrationConnection.id == connection_id,
                IntegrationConnection.organization_id == organization_id,
            )
        )
        if connection is None:
            raise HTTPException(status_code=404, detail="Integration connection not found.")
        connection.status = "revoked"
        connection.updated_at = utcnow()
        self._audit(
            organization_id,
            user.id,
            "integration.revoked",
            detail={"connection_id": connection.id, "provider": connection.provider},
        )
        self.session.commit()
        self.session.refresh(connection)
        return connection

    def create_imported_contract(
        self,
        organization_id: str,
        user: User,
        *,
        provider: str,
        source_type: str,
        original_name: str,
        content_type: str,
        data: bytes,
        title: str,
        counterparty: str = "",
        contract_type: str = "Unknown",
        connection_id: str | None = None,
        external_id: str = "",
        source_url: str = "",
        metadata: dict[str, Any] | None = None,
        retain_document: bool = False,
        retain_source_text: bool = False,
        retention_days: int = 30,
    ) -> tuple[IntegrationImport, Contract, DocumentAsset, ProcessingJob]:
        if provider not in INTEGRATION_PROVIDERS:
            raise HTTPException(status_code=422, detail="Unknown integration provider.")
        connection = None
        if connection_id:
            connection = self.session.scalar(
                select(IntegrationConnection).where(
                    IntegrationConnection.id == connection_id,
                    IntegrationConnection.organization_id == organization_id,
                    IntegrationConnection.provider == provider,
                    IntegrationConnection.status == "active",
                )
            )
            if connection is None:
                raise HTTPException(status_code=422, detail="Choose an active connection for this provider.")
        stable_external_id = external_id.strip() or hashlib.sha256(data).hexdigest()
        existing = self.session.scalar(
            select(IntegrationImport).where(
                IntegrationImport.organization_id == organization_id,
                IntegrationImport.provider == provider,
                IntegrationImport.external_id == stable_external_id,
            )
        )
        if existing is not None:
            raise HTTPException(status_code=409, detail="This source document has already been imported.")
        contract, asset, job = self.create_contract(
            organization_id,
            user,
            original_name=original_name,
            content_type=content_type,
            data=data,
            title=title,
            counterparty=counterparty,
            contract_type=contract_type,
            review_context={
                "party_role": "Not sure / general review",
                "jurisdiction": "",
                "goal": f"Imported from {provider.replace('_', ' ')}",
                "risk_tolerance": "Balanced",
                "intake_provider": provider,
                "intake_source_type": source_type,
            },
            retain_document=retain_document,
            retain_source_text=retain_source_text,
            retention_days=retention_days,
        )
        import_record = IntegrationImport(
            organization_id=organization_id,
            connection_id=connection.id if connection else None,
            contract_id=contract.id,
            imported_by_user_id=user.id,
            provider=provider,
            source_type=source_type,
            external_id=stable_external_id,
            source_url=source_url.strip()[:1024],
            title=(title or contract.title).strip()[:512],
            original_name=safe_filename(original_name),
            content_type=content_type or "application/octet-stream",
            size_bytes=len(data),
            sha256=asset.sha256,
            status="queued",
            metadata_json=json_dump(metadata or {}),
        )
        self.session.add(import_record)
        self._audit(
            organization_id,
            user.id,
            "integration.import_created",
            contract.id,
            {
                "import_id": import_record.id,
                "provider": provider,
                "source_type": source_type,
                "external_id": stable_external_id,
            },
        )
        self._notify_import(contract, provider)
        self.session.commit()
        self.session.refresh(import_record)
        self.session.refresh(contract)
        self.session.refresh(asset)
        self.session.refresh(job)
        return import_record, contract, asset, job

    def list_integration_imports(
        self,
        organization_id: str,
        user: User,
        provider: str | None = None,
    ) -> list[IntegrationImport]:
        self.membership(organization_id, user)
        query = select(IntegrationImport).where(IntegrationImport.organization_id == organization_id)
        if provider:
            if provider not in INTEGRATION_PROVIDERS:
                raise HTTPException(status_code=422, detail="Unknown integration provider.")
            query = query.where(IntegrationImport.provider == provider)
        return list(
            self.session.scalars(query.order_by(IntegrationImport.created_at.desc()).limit(100)).all()
        )

    def _notify_import(self, contract: Contract, provider: str) -> None:
        member_ids = self.session.scalars(
            select(Membership.user_id).where(Membership.organization_id == contract.organization_id)
        ).all()
        for user_id in set(member_ids):
            self._notify(
                contract.organization_id,
                user_id,
                contract.id,
                "contract_imported",
                "Contract imported",
                f"{contract.title} was imported from {provider.replace('_', ' ')}.",
                f"/contracts/{contract.id}",
            )

    def create_api_key(self, organization_id: str, user: User, payload: dict[str, Any]) -> ApiKeyCreatedResponse:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can manage API keys.",
        )
        scopes = [scope for scope in dict.fromkeys(payload.get("scopes", [])) if scope in API_KEY_SCOPES]
        if not scopes:
            raise HTTPException(status_code=422, detail="Choose at least one supported API key scope.")
        token = f"ll_live_{secrets.token_urlsafe(32)}"
        api_key = PublicApiKey(
            organization_id=organization_id,
            created_by_user_id=user.id,
            name=payload["name"].strip(),
            key_prefix=token[:16],
            key_hash=hashlib.sha256(token.encode("utf-8")).hexdigest(),
            scopes_json=json_dump(scopes),
        )
        self.session.add(api_key)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "api_key.created",
            detail={"api_key_id": api_key.id, "name": api_key.name, "scopes": scopes},
        )
        self.session.commit()
        self.session.refresh(api_key)
        return ApiKeyCreatedResponse(api_key=self.api_key_response(api_key), token=token)

    def list_api_keys(self, organization_id: str, user: User) -> list[PublicApiKey]:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can view API keys.",
        )
        return list(
            self.session.scalars(
                select(PublicApiKey)
                .where(PublicApiKey.organization_id == organization_id)
                .order_by(PublicApiKey.created_at.desc())
            ).all()
        )

    def revoke_api_key(self, organization_id: str, api_key_id: str, user: User) -> None:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can revoke API keys.",
        )
        api_key = self.session.scalar(
            select(PublicApiKey).where(
                PublicApiKey.id == api_key_id,
                PublicApiKey.organization_id == organization_id,
            )
        )
        if api_key is None:
            raise HTTPException(status_code=404, detail="API key not found.")
        api_key.revoked_at = utcnow()
        self._audit(organization_id, user.id, "api_key.revoked", detail={"api_key_id": api_key.id})
        self.session.commit()

    def authenticate_api_key(self, token: str, required_scope: str) -> tuple[PublicApiKey, User]:
        key_hash = hashlib.sha256(token.encode("utf-8")).hexdigest()
        api_key = self.session.scalar(select(PublicApiKey).where(PublicApiKey.key_hash == key_hash))
        if api_key is None or api_key.revoked_at is not None:
            raise HTTPException(status_code=401, detail="A valid Lenslayer API key is required.")
        scopes = set(json_load(api_key.scopes_json, []))
        if required_scope not in scopes:
            raise HTTPException(status_code=403, detail="The API key does not allow this operation.")
        api_key.last_used_at = utcnow()
        self.session.commit()
        return api_key, api_key.created_by

    def create_webhook_subscription(
        self,
        organization_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> tuple[WebhookSubscription, str]:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can manage webhooks.",
        )
        events = [event for event in dict.fromkeys(payload.get("events", [])) if event in WEBHOOK_EVENTS]
        if not events:
            raise HTTPException(status_code=422, detail="Choose at least one supported webhook event.")
        signing_secret = f"whsec_{secrets.token_urlsafe(32)}"
        subscription = WebhookSubscription(
            organization_id=organization_id,
            created_by_user_id=user.id,
            target_url=payload["target_url"].strip(),
            description=payload.get("description", "").strip(),
            events_json=json_dump(events),
            signing_secret_hash=hashlib.sha256(signing_secret.encode("utf-8")).hexdigest(),
            secret_prefix=signing_secret[:12],
            status="active",
        )
        self.session.add(subscription)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "webhook.created",
            detail={"webhook_id": subscription.id, "events": events},
        )
        self.session.commit()
        self.session.refresh(subscription)
        return subscription, signing_secret

    def list_webhook_subscriptions(self, organization_id: str, user: User) -> list[WebhookSubscription]:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can view webhooks.",
        )
        return list(
            self.session.scalars(
                select(WebhookSubscription)
                .where(WebhookSubscription.organization_id == organization_id)
                .order_by(WebhookSubscription.created_at.desc())
            ).all()
        )

    def revoke_webhook_subscription(self, organization_id: str, webhook_id: str, user: User) -> None:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can revoke webhooks.",
        )
        webhook = self.session.scalar(
            select(WebhookSubscription).where(
                WebhookSubscription.id == webhook_id,
                WebhookSubscription.organization_id == organization_id,
            )
        )
        if webhook is None:
            raise HTTPException(status_code=404, detail="Webhook not found.")
        webhook.status = "revoked"
        webhook.updated_at = utcnow()
        self._audit(organization_id, user.id, "webhook.revoked", detail={"webhook_id": webhook.id})
        self.session.commit()

    def list_webhook_deliveries(
        self,
        organization_id: str,
        user: User,
        webhook_id: str | None = None,
    ) -> list[WebhookDelivery]:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can view webhook deliveries.",
        )
        query = select(WebhookDelivery).where(WebhookDelivery.organization_id == organization_id)
        if webhook_id:
            query = query.where(WebhookDelivery.subscription_id == webhook_id)
        return list(
            self.session.scalars(query.order_by(WebhookDelivery.created_at.desc()).limit(100)).all()
        )

    def _enqueue_webhooks(
        self,
        organization_id: str,
        event_type: str,
        contract_id: str | None,
        payload: dict[str, Any],
    ) -> None:
        subscriptions = self.session.scalars(
            select(WebhookSubscription).where(
                WebhookSubscription.organization_id == organization_id,
                WebhookSubscription.status == "active",
            )
        ).all()
        for subscription in subscriptions:
            if event_type not in set(json_load(subscription.events_json, [])):
                continue
            self.session.add(
                WebhookDelivery(
                    organization_id=organization_id,
                    subscription_id=subscription.id,
                    contract_id=contract_id,
                    event_type=event_type,
                    payload_json=json_dump({"event": event_type, "data": payload}),
                    status="pending",
                )
            )

    def list_contracts(self, organization_id: str, user: User) -> list[Contract]:
        self.membership(organization_id, user)
        return list(
            self.session.scalars(
                select(Contract)
                .where(Contract.organization_id == organization_id)
                .order_by(Contract.updated_at.desc())
            ).all()
        )

    def get_contract(self, organization_id: str, contract_id: str, user: User) -> Contract:
        self.membership(organization_id, user)
        contract = self.session.scalar(
            select(Contract).where(
                Contract.id == contract_id,
                Contract.organization_id == organization_id,
            )
        )
        if contract is None:
            raise HTTPException(status_code=404, detail="Contract not found.")
        return contract

    def get_review(self, organization_id: str, contract_id: str, user: User) -> ContractReview:
        contract = self.get_contract(organization_id, contract_id, user)
        if contract.review is None:
            raise HTTPException(status_code=404, detail="The contract review is not ready.")
        return contract.review

    def list_contract_versions(self, organization_id: str, contract_id: str, user: User) -> list[ContractVersion]:
        self.get_contract(organization_id, contract_id, user)
        return list(
            self.session.scalars(
                select(ContractVersion)
                .where(
                    ContractVersion.organization_id == organization_id,
                    ContractVersion.contract_id == contract_id,
                )
                .order_by(ContractVersion.version_number.asc())
            ).all()
        )

    def create_contract_version(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        *,
        original_name: str,
        content_type: str,
        data: bytes,
        label: str,
        notes: str,
    ) -> ContractVersion:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot upload revised documents.",
        )
        contract = self.get_contract(organization_id, contract_id, user)
        filename = safe_filename(original_name)
        suffix = Path(filename).suffix.lower()
        if suffix not in self.settings.allowed_extension_set:
            raise HTTPException(status_code=415, detail="Supported file types are PDF, DOCX, and TXT.")
        if not data:
            raise HTTPException(status_code=400, detail="The uploaded document is empty.")
        if len(data) > self.settings.max_upload_bytes:
            raise HTTPException(status_code=413, detail="The uploaded document exceeds the 25 MB limit.")
        scan_upload(self.settings, filename, data)

        previous = self.session.scalar(
            select(ContractVersion)
            .where(
                ContractVersion.organization_id == organization_id,
                ContractVersion.contract_id == contract_id,
            )
            .order_by(ContractVersion.version_number.desc())
            .limit(1)
        )
        version_number = (previous.version_number + 1) if previous else 1
        asset_id = str(uuid4())
        storage_key = f"{organization_id}/{contract.id}/{asset_id}{suffix}"
        sha256 = hashlib.sha256(data).hexdigest()
        new_text = self._extract_uploaded_text(filename, data)
        old_text = self._version_text(previous) if previous else ""
        comparison = self._compare_version_texts(old_text, new_text, previous.id if previous else None)
        asset_status = "available" if contract.retain_document else "deleted"
        asset = DocumentAsset(
            id=asset_id,
            organization_id=organization_id,
            contract_id=contract.id,
            storage_key=storage_key,
            original_name=filename,
            content_type=content_type or "application/octet-stream",
            size_bytes=len(data),
            sha256=sha256,
            status=asset_status,
        )
        version = ContractVersion(
            organization_id=organization_id,
            contract_id=contract.id,
            document_asset_id=asset.id,
            uploaded_by_user_id=user.id,
            version_number=version_number,
            label=(label or f"Version {version_number}").strip()[:255],
            notes=notes.strip(),
            source_name=filename,
            sha256=sha256,
            size_bytes=len(data),
            comparison_json=json_dump(comparison),
            extracted_text=new_text if contract.retain_source_text else None,
        )
        self.session.add_all((asset, version))
        self._audit(
            organization_id,
            user.id,
            "contract.version_uploaded",
            contract.id,
            {
                "version_number": version_number,
                "source_name": filename,
                "added_count": comparison["added_count"],
                "removed_count": comparison["removed_count"],
            },
        )
        try:
            self.object_store.put(storage_key, data, asset.content_type)
            if not contract.retain_document:
                self.object_store.delete(storage_key)
            contract.updated_at = utcnow()
            self.session.commit()
        except Exception:
            self.session.rollback()
            self.object_store.delete(storage_key)
            raise
        self.session.refresh(version)
        return version

    def _extract_uploaded_text(self, filename: str, data: bytes) -> str:
        suffix = Path(filename).suffix.lower()
        if suffix == ".txt":
            return data.decode("utf-8", errors="replace")
        temp_path: str | None = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as handle:
                handle.write(data)
                temp_path = handle.name
            from analyzer import parse_document

            text, _, _ = parse_document(temp_path)
            return text
        finally:
            if temp_path:
                try:
                    os.unlink(temp_path)
                except OSError:
                    pass

    def _version_text(self, version: ContractVersion | None) -> str:
        if version is None:
            return ""
        if version.extracted_text:
            return version.extracted_text
        if version.document_asset and version.document_asset.status == "available":
            try:
                data = self.object_store.get(version.document_asset.storage_key)
            except Exception:
                return ""
            return self._extract_uploaded_text(version.document_asset.original_name, data)
        return ""

    @staticmethod
    def _compare_version_texts(old_text: str, new_text: str, previous_id: str | None) -> dict[str, Any]:
        old_lines = [line.strip() for line in old_text.splitlines() if line.strip()]
        new_lines = [line.strip() for line in new_text.splitlines() if line.strip()]
        added: list[str] = []
        removed: list[str] = []
        for line in difflib.ndiff(old_lines, new_lines):
            if line.startswith("+ "):
                added.append(line[2:])
            elif line.startswith("- "):
                removed.append(line[2:])
        if not old_lines:
            summary = "No retained prior text was available for comparison."
        elif not added and not removed:
            summary = "No substantive line-level changes detected."
        else:
            summary = f"{len(added)} added and {len(removed)} removed line-level change{'s' if len(added) + len(removed) != 1 else ''} detected."
        return {
            "compared_to_version_id": previous_id,
            "added": added[:50],
            "removed": removed[:50],
            "changed_summary": summary,
            "added_count": len(added),
            "removed_count": len(removed),
        }

    @staticmethod
    def contract_version_response(version: ContractVersion) -> ContractVersionResponse:
        comparison = json_load(version.comparison_json, {})
        return ContractVersionResponse(
            id=version.id,
            contract_id=version.contract_id,
            document_asset_id=version.document_asset_id,
            version_number=version.version_number,
            label=version.label,
            notes=version.notes,
            source_name=version.source_name,
            sha256=version.sha256,
            size_bytes=version.size_bytes,
            comparison={
                "compared_to_version_id": comparison.get("compared_to_version_id"),
                "added": comparison.get("added", []),
                "removed": comparison.get("removed", []),
                "changed_summary": comparison.get("changed_summary", ""),
                "added_count": comparison.get("added_count", 0),
                "removed_count": comparison.get("removed_count", 0),
            },
            uploaded_by_user_id=version.uploaded_by_user_id,
            uploaded_by_name=version.uploaded_by.display_name if version.uploaded_by else "System",
            created_at=version.created_at,
        )

    def list_negotiation_items(self, organization_id: str, contract_id: str, user: User) -> list[NegotiationItem]:
        self.get_contract(organization_id, contract_id, user)
        return list(
            self.session.scalars(
                select(NegotiationItem)
                .where(
                    NegotiationItem.organization_id == organization_id,
                    NegotiationItem.contract_id == contract_id,
                )
                .order_by(NegotiationItem.created_at.asc())
            ).all()
        )

    def create_negotiation_item(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> NegotiationItem:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot change the negotiation checklist.",
        )
        self.get_contract(organization_id, contract_id, user)
        item = NegotiationItem(
            organization_id=organization_id,
            contract_id=contract_id,
            created_by_user_id=user.id,
            title=payload["title"].strip(),
            description=payload.get("description", "").strip(),
            category=payload.get("category", "change"),
            priority=payload.get("priority", "normal"),
            status=payload.get("status", "proposed"),
            our_position=payload.get("our_position", "").strip(),
            counterparty_position=payload.get("counterparty_position", "").strip(),
            source_reference_json=json_dump(payload.get("source_reference", {})),
            resolved_at=utcnow() if payload.get("status") in {"accepted", "rejected", "resolved"} else None,
        )
        self.session.add(item)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "negotiation.item_created",
            contract_id,
            {"negotiation_item_id": item.id, "status": item.status, "priority": item.priority},
        )
        self.session.commit()
        self.session.refresh(item)
        return item

    def update_negotiation_item(
        self,
        organization_id: str,
        contract_id: str,
        item_id: str,
        user: User,
        changes: dict[str, Any],
    ) -> NegotiationItem:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot change the negotiation checklist.",
        )
        item = self.session.scalar(
            select(NegotiationItem).where(
                NegotiationItem.id == item_id,
                NegotiationItem.organization_id == organization_id,
                NegotiationItem.contract_id == contract_id,
            )
        )
        if item is None:
            raise HTTPException(status_code=404, detail="Negotiation item not found.")
        for field in ("title", "description", "category", "priority", "status", "our_position", "counterparty_position"):
            if field in changes and changes[field] is not None:
                value = changes[field]
                if isinstance(value, str):
                    value = value.strip()
                setattr(item, field, value)
        if "source_reference" in changes and changes["source_reference"] is not None:
            item.source_reference_json = json_dump(changes["source_reference"])
        if "status" in changes and changes["status"] is not None:
            item.resolved_at = utcnow() if item.status in {"accepted", "rejected", "resolved"} else None
        self._audit(
            organization_id,
            user.id,
            "negotiation.item_updated",
            contract_id,
            {"negotiation_item_id": item.id, "changed_fields": sorted(changes)},
        )
        self.session.commit()
        self.session.refresh(item)
        return item

    @staticmethod
    def negotiation_item_response(item: NegotiationItem) -> NegotiationItemResponse:
        return NegotiationItemResponse(
            id=item.id,
            contract_id=item.contract_id,
            title=item.title,
            description=item.description,
            category=item.category,
            priority=item.priority,
            status=item.status,
            our_position=item.our_position,
            counterparty_position=item.counterparty_position,
            source_reference=json_load(item.source_reference_json, {}),
            created_by_user_id=item.created_by_user_id,
            created_by_name=item.created_by.display_name,
            resolved_at=item.resolved_at,
            created_at=item.created_at,
            updated_at=item.updated_at,
        )

    def list_counterparty_responses(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
    ) -> list[CounterpartyResponse]:
        self.get_contract(organization_id, contract_id, user)
        return list(
            self.session.scalars(
                select(CounterpartyResponse)
                .where(
                    CounterpartyResponse.organization_id == organization_id,
                    CounterpartyResponse.contract_id == contract_id,
                )
                .order_by(CounterpartyResponse.created_at.desc())
            ).all()
        )

    def create_counterparty_response(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> CounterpartyResponse:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot record counterparty responses.",
        )
        self.get_contract(organization_id, contract_id, user)
        version_id = payload.get("contract_version_id")
        if version_id:
            version = self.session.scalar(
                select(ContractVersion).where(
                    ContractVersion.id == version_id,
                    ContractVersion.organization_id == organization_id,
                    ContractVersion.contract_id == contract_id,
                )
            )
            if version is None:
                raise HTTPException(status_code=422, detail="Choose a version from this contract.")
        related_item_ids = list(dict.fromkeys(payload.get("related_item_ids", [])))
        if related_item_ids:
            existing_ids = set(
                self.session.scalars(
                    select(NegotiationItem.id).where(
                        NegotiationItem.organization_id == organization_id,
                        NegotiationItem.contract_id == contract_id,
                        NegotiationItem.id.in_(related_item_ids),
                    )
                ).all()
            )
            if existing_ids != set(related_item_ids):
                raise HTTPException(status_code=422, detail="Choose checklist items from this contract.")
        response = CounterpartyResponse(
            organization_id=organization_id,
            contract_id=contract_id,
            contract_version_id=version_id,
            recorded_by_user_id=user.id,
            responder_name=payload.get("responder_name", "").strip()[:255],
            channel=payload.get("channel", "email").strip()[:64] or "email",
            body=payload["body"].strip(),
            related_item_ids_json=json_dump(related_item_ids),
        )
        self.session.add(response)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "negotiation.counterparty_response_recorded",
            contract_id,
            {"response_id": response.id, "related_item_ids": related_item_ids, "contract_version_id": version_id},
        )
        self.session.commit()
        self.session.refresh(response)
        return response

    @staticmethod
    def counterparty_response_response(response: CounterpartyResponse) -> CounterpartyResponseResponse:
        return CounterpartyResponseResponse(
            id=response.id,
            contract_id=response.contract_id,
            contract_version_id=response.contract_version_id,
            recorded_by_user_id=response.recorded_by_user_id,
            recorded_by_name=response.recorded_by.display_name,
            responder_name=response.responder_name,
            channel=response.channel,
            body=response.body,
            related_item_ids=json_load(response.related_item_ids_json, []),
            created_at=response.created_at,
        )

    def negotiation_summary(self, organization_id: str, contract_id: str, user: User) -> NegotiationSummaryResponse:
        self.get_contract(organization_id, contract_id, user)
        versions = self.list_contract_versions(organization_id, contract_id, user)
        items = self.list_negotiation_items(organization_id, contract_id, user)
        responses = self.list_counterparty_responses(organization_id, contract_id, user)
        accepted = [item for item in items if item.status == "accepted"]
        rejected = [item for item in items if item.status == "rejected"]
        unresolved = [item for item in items if item.status in {"proposed", "unresolved"}]
        latest = versions[-1] if versions else None
        parts = [
            f"{len(versions)} document version{'s' if len(versions) != 1 else ''} retained in the history.",
            f"{len(accepted)} accepted change{'s' if len(accepted) != 1 else ''}, {len(rejected)} rejected change{'s' if len(rejected) != 1 else ''}, and {len(unresolved)} unresolved point{'s' if len(unresolved) != 1 else ''}.",
            f"{len(responses)} counterparty response{'s' if len(responses) != 1 else ''} recorded.",
        ]
        if latest:
            parts.append(f"Latest version: {latest.label or latest.source_name} ({latest.comparison_json and json_load(latest.comparison_json, {}).get('changed_summary', '')}).")
        return NegotiationSummaryResponse(
            contract_id=contract_id,
            latest_version=self.contract_version_response(latest) if latest else None,
            version_count=len(versions),
            checklist_count=len(items),
            accepted_changes=[self.negotiation_item_response(item) for item in accepted],
            rejected_changes=[self.negotiation_item_response(item) for item in rejected],
            unresolved_points=[self.negotiation_item_response(item) for item in unresolved],
            counterparty_response_count=len(responses),
            final_summary=" ".join(part for part in parts if part),
        )

    def deal_passport(self, organization_id: str, contract_id: str, user: User) -> DealPassportResponse:
        contract = self.get_contract(organization_id, contract_id, user)
        review = self.get_review(organization_id, contract_id, user)
        analysis = json_load(review.analysis_json, {})
        versions = self.list_contract_versions(organization_id, contract_id, user)
        negotiation = self.negotiation_summary(organization_id, contract_id, user)
        approvals = list(self.session.scalars(select(ApprovalRequest).where(
            ApprovalRequest.organization_id == organization_id,
            ApprovalRequest.contract_id == contract_id,
        ).order_by(ApprovalRequest.created_at.desc())).all())
        tasks = list(self.session.scalars(select(WorkflowTask).where(
            WorkflowTask.organization_id == organization_id,
            WorkflowTask.contract_id == contract_id,
            WorkflowTask.status.in_(("open", "in_progress")),
        ).order_by(WorkflowTask.due_at.asc())).all())
        lifecycle = list(self.session.scalars(select(LifecycleItem).where(
            LifecycleItem.organization_id == organization_id,
            LifecycleItem.contract_id == contract_id,
            LifecycleItem.status == "active",
        ).order_by(LifecycleItem.due_at.asc())).all())
        risks = sorted(
            [item for item in analysis.get("risk_assessment", []) if isinstance(item, dict)],
            key=lambda item: {"high": 0, "medium": 1, "low": 2}.get(str(item.get("risk_level", "")).casefold(), 3),
        )[:3]
        reasons: list[str] = []
        if contract.status != "ready":
            reasons.append("The automated review is not ready.")
        high_risks = [item for item in risks if str(item.get("risk_level", "")).casefold() == "high"]
        if high_risks:
            reasons.append(f"{len(high_risks)} high-priority risk finding{'s' if len(high_risks) != 1 else ''} need a decision.")
        if negotiation.unresolved_points:
            reasons.append(f"{len(negotiation.unresolved_points)} negotiation point{'s' if len(negotiation.unresolved_points) != 1 else ''} remain unresolved.")
        pending_approvals = [item for item in approvals if item.status in {"pending", "changes_requested"}]
        if pending_approvals:
            reasons.append(f"{len(pending_approvals)} approval request{'s' if len(pending_approvals) != 1 else ''} remain open.")
        readiness = "blocked" if contract.status != "ready" or pending_approvals else "needs_attention" if reasons or tasks else "ready"
        return DealPassportResponse(
            contract_id=contract.id,
            title=contract.title,
            counterparty=contract.counterparty,
            contract_type=contract.contract_type,
            readiness=readiness,
            readiness_reasons=reasons,
            executive_summary=str(analysis.get("executive_summary") or ""),
            overall_attention=str(analysis.get("overall_attention") or ""),
            top_risks=risks,
            versions=[self.contract_version_response(item) for item in versions],
            negotiation=negotiation,
            approvals=[{
                "id": item.id,
                "title": item.title,
                "status": item.status,
                "assigned_to": item.assigned_to.display_name if item.assigned_to else "",
                "due_at": item.due_at,
            } for item in approvals],
            open_actions=[{
                "id": item.id,
                "title": item.title,
                "priority": item.priority,
                "status": item.status,
                "owner": item.assigned_to_user.display_name if item.assigned_to_user else "",
                "due_at": item.due_at,
            } for item in tasks],
            key_dates=[{
                "id": item.id,
                "kind": item.kind,
                "title": item.title,
                "due_at": item.due_at,
                "owner": item.owner.display_name if item.owner else "",
            } for item in lifecycle],
            generated_at=utcnow(),
        )

    def redline_export(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
    ) -> tuple[bytes, str]:
        contract = self.get_contract(organization_id, contract_id, user)
        review = self.get_review(organization_id, contract_id, user)
        version = self.session.scalar(select(ContractVersion).where(
            ContractVersion.organization_id == organization_id,
            ContractVersion.contract_id == contract_id,
        ).order_by(ContractVersion.version_number.desc()).limit(1))
        asset = version.document_asset if version else None
        if not asset or asset.status != "available" or Path(asset.original_name).suffix.lower() != ".docx":
            raise HTTPException(
                status_code=409,
                detail="Tracked-change redlining requires a retained DOCX version. Upload a DOCX with document retention enabled.",
            )
        findings = json_load(review.analysis_json, {}).get("risk_assessment", [])
        from .redline import build_redline
        content, change_count = build_redline(self.object_store.get(asset.storage_key), findings)
        if not change_count:
            raise HTTPException(status_code=409, detail="This review has no suggested replacement language to redline.")
        self._audit(organization_id, user.id, "review.redline_exported", contract_id, {
            "version_number": version.version_number,
            "suggested_change_count": change_count,
        })
        self.session.commit()
        return content, f"{safe_filename(Path(contract.source_name).stem)}-lenslayer-redline.docx"

    def answer_contract_question(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        question: str,
    ) -> tuple[str, list[dict[str, str]], str]:
        review = self.get_review(organization_id, contract_id, user)
        source_text = (review.source_text or "").strip()
        if not source_text:
            raise HTTPException(
                status_code=409,
                detail="Contract Q&A requires retained source text. Upload the contract again with source-text retention enabled.",
            )
        tokens = {
            token
            for token in re.findall(r"[a-z0-9]{3,}", question.casefold())
            if token not in {"about", "what", "when", "where", "which", "that", "this", "with", "from", "does", "have"}
        }
        blocks = [item.strip() for item in re.split(r"\n{2,}", source_text) if item.strip()]
        scored: list[tuple[int, int, str]] = []
        for index, block in enumerate(blocks):
            haystack = block.casefold()
            score = sum(haystack.count(token) for token in tokens)
            scored.append((score, -index, block))
        selected = [item[2] for item in sorted(scored, reverse=True)[:4] if item[0] > 0]
        if not selected:
            selected = blocks[:3]
        sources = [
            {
                "label": f"Source {index}",
                "location": self._source_location(block, index),
                "excerpt": " ".join(block.split())[:900],
            }
            for index, block in enumerate(selected, start=1)
        ]
        if not sources:
            raise HTTPException(status_code=409, detail="The retained source text is empty.")

        context = "\n\n".join(
            f"[{source['label']}] {source['location']}\n{source['excerpt']}" for source in sources
        )
        if os.environ.get("GROQ_API_KEY"):
            from analyzer import get_llm, review_context_text

            contract = self.get_contract(organization_id, contract_id, user)
            prompt = (
                "Answer this contract question using only the evidence below. Cite sources inline as "
                "[Source 1]. If the evidence does not establish the answer, say so. Use plain language "
                "and distinguish document facts from suggested next steps. This is first-pass education, "
                "not legal advice.\n\n"
                f"REVIEW CONTEXT\n{review_context_text(json_load(contract.review_context_json, {}))}\n\n"
                f"QUESTION\n{question.strip()}\n\nEVIDENCE\n{context}"
            )
            response = get_llm().invoke(prompt)
            answer = str(response.content).strip()
            generated_by = "model"
        else:
            answer = (
                "The most relevant retained excerpts are shown below. A model answer is unavailable in this "
                "environment, so confirm the wording directly in the document before acting."
            )
            generated_by = "extractive"
        self._audit(
            organization_id,
            user.id,
            "review.question_answered",
            contract_id,
            {"question": question.strip()[:240], "source_count": len(sources), "generated_by": generated_by},
        )
        self.session.commit()
        return answer, sources, generated_by

    @staticmethod
    def _source_location(block: str, index: int) -> str:
        match = re.search(r"\[(PAGE \d+|L\d+)\]", block, flags=re.IGNORECASE)
        return match.group(1).title() if match else f"Retained excerpt {index}"

    def contract_export(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        export_format: str,
    ) -> tuple[bytes, str, str]:
        contract = self.get_contract(organization_id, contract_id, user)
        review = self.get_review(organization_id, contract_id, user)
        analysis = json_load(review.analysis_json, {})
        quality = json_load(review.quality_json, {})
        context = json_load(contract.review_context_json, {})
        from export_utils import (
            build_csv,
            build_docx_report,
            build_json_report,
            build_markdown_report,
            build_pdf_report,
        )

        safe_stem = safe_filename(Path(contract.source_name).stem)
        if export_format == "pdf":
            content = build_pdf_report(analysis, contract.source_name, context)
            media_type, suffix = "application/pdf", "pdf"
        elif export_format == "docx":
            content = build_docx_report(analysis, contract.source_name, context)
            media_type, suffix = "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
        elif export_format == "csv":
            rows: list[dict[str, Any]] = []
            for category, items in (
                ("risk", analysis.get("risk_assessment", [])),
                ("protection_gap", analysis.get("missing_protections", [])),
                ("obligation", analysis.get("obligations", [])),
                ("deadline", analysis.get("deadlines", [])),
                ("payment", analysis.get("payments", [])),
                ("negotiation", analysis.get("negotiation_priorities", [])),
            ):
                for item in items or []:
                    rows.append({"category": category, **(item if isinstance(item, dict) else {"detail": item})})
            content = build_csv(rows).encode("utf-8")
            media_type, suffix = "text/csv; charset=utf-8", "csv"
        elif export_format == "md":
            content = build_markdown_report(analysis, contract.source_name, context).encode("utf-8")
            media_type, suffix = "text/markdown; charset=utf-8", "md"
        elif export_format == "json":
            content = build_json_report(analysis, context, quality).encode("utf-8")
            media_type, suffix = "application/json; charset=utf-8", "json"
        else:
            raise HTTPException(status_code=422, detail="Choose PDF, DOCX, CSV, Markdown, or JSON.")
        self._audit(
            organization_id,
            user.id,
            "review.exported",
            contract_id,
            {"format": export_format},
        )
        self.session.commit()
        return content, media_type, f"{safe_stem}-lenslayer-review.{suffix}"

    def list_contract_comments(self, organization_id: str, contract_id: str, user: User) -> list[ContractComment]:
        self.get_contract(organization_id, contract_id, user)
        return list(
            self.session.scalars(
                select(ContractComment)
                .where(
                    ContractComment.organization_id == organization_id,
                    ContractComment.contract_id == contract_id,
                )
                .order_by(ContractComment.created_at.asc())
            ).all()
        )

    def create_contract_comment(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        body: str,
        mentioned_user_ids: list[str],
    ) -> ContractComment:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot add comments.",
        )
        contract = self.get_contract(organization_id, contract_id, user)
        valid_mentions: list[str] = []
        for mentioned_user_id in dict.fromkeys(mentioned_user_ids):
            if mentioned_user_id == user.id:
                continue
            member = self.session.scalar(
                select(Membership).where(
                    Membership.organization_id == organization_id,
                    Membership.user_id == mentioned_user_id,
                )
            )
            if member is None:
                raise HTTPException(status_code=422, detail="A mentioned user is not a workspace member.")
            valid_mentions.append(mentioned_user_id)
        comment = ContractComment(
            organization_id=organization_id,
            contract_id=contract_id,
            author_user_id=user.id,
            body=body.strip(),
            mentions_json=json_dump(valid_mentions),
        )
        self.session.add(comment)
        self.session.flush()
        for mentioned_user_id in valid_mentions:
            self._notify(
                organization_id,
                mentioned_user_id,
                contract_id,
                "mention",
                f"{user.display_name or user.email} mentioned you",
                f"New comment on {contract.title}.",
                f"/contracts/{contract_id}?tab=collaboration",
            )
        self._audit(
            organization_id,
            user.id,
            "comment.created",
            contract_id,
            {"comment_id": comment.id, "mentioned_user_ids": valid_mentions},
        )
        self.session.commit()
        self.session.refresh(comment)
        return comment

    @staticmethod
    def contract_comment_response(comment: ContractComment) -> ContractCommentResponse:
        return ContractCommentResponse(
            id=comment.id,
            body=comment.body,
            mentioned_user_ids=json_load(comment.mentions_json, []),
            author_user_id=comment.author_user_id,
            author_name=comment.author.display_name,
            author_email=comment.author.email,
            created_at=comment.created_at,
            updated_at=comment.updated_at,
        )

    def list_contract_decisions(self, organization_id: str, contract_id: str, user: User) -> list[ContractDecision]:
        self.get_contract(organization_id, contract_id, user)
        return list(
            self.session.scalars(
                select(ContractDecision)
                .where(
                    ContractDecision.organization_id == organization_id,
                    ContractDecision.contract_id == contract_id,
                )
                .order_by(ContractDecision.created_at.asc())
            ).all()
        )

    def create_contract_decision(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> ContractDecision:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot record review decisions.",
        )
        self.get_contract(organization_id, contract_id, user)
        if payload["decision"] not in CONTRACT_DECISIONS:
            raise HTTPException(status_code=422, detail="Choose accept, change, escalate, or resolve.")
        decision = ContractDecision(
            organization_id=organization_id,
            contract_id=contract_id,
            reviewer_user_id=user.id,
            decision=payload["decision"],
            subject=payload["subject"].strip(),
            rationale=payload["rationale"].strip(),
            source_reference_json=json_dump(payload.get("source_reference", {})),
        )
        self.session.add(decision)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "contract.decision_recorded",
            contract_id,
            {"decision_id": decision.id, "decision": decision.decision, "subject": decision.subject},
        )
        self.session.commit()
        self.session.refresh(decision)
        return decision

    @staticmethod
    def contract_decision_response(decision: ContractDecision) -> ContractDecisionResponse:
        return ContractDecisionResponse(
            id=decision.id,
            decision=decision.decision,
            subject=decision.subject,
            rationale=decision.rationale,
            source_reference=json_load(decision.source_reference_json, {}),
            reviewer_user_id=decision.reviewer_user_id,
            reviewer_name=decision.reviewer.display_name,
            reviewer_email=decision.reviewer.email,
            created_at=decision.created_at,
        )

    def list_approval_requests(self, organization_id: str, contract_id: str, user: User) -> list[ApprovalRequest]:
        self.get_contract(organization_id, contract_id, user)
        return list(
            self.session.scalars(
                select(ApprovalRequest)
                .where(
                    ApprovalRequest.organization_id == organization_id,
                    ApprovalRequest.contract_id == contract_id,
                )
                .order_by(ApprovalRequest.created_at.desc())
            ).all()
        )

    def create_approval_request(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> ApprovalRequest:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot request approval.",
        )
        contract = self.get_contract(organization_id, contract_id, user)
        assignee = self._task_assignee(organization_id, payload.get("assigned_to_user_id"))
        request = ApprovalRequest(
            organization_id=organization_id,
            contract_id=contract_id,
            requested_by_user_id=user.id,
            assigned_to_user_id=assignee.id if assignee else None,
            title=payload["title"].strip(),
            note=payload.get("note", "").strip(),
            conditions_json=json_dump([item.strip() for item in payload.get("conditions", []) if item.strip()]),
            due_at=payload.get("due_at"),
        )
        self.session.add(request)
        self.session.flush()
        recipients = [assignee.id] if assignee else [
            item.user_id for item in self.session.scalars(
                select(Membership).where(
                    Membership.organization_id == organization_id,
                    Membership.role.in_(["owner", "admin"]),
                )
            ).all()
        ]
        for recipient_id in set(recipients):
            if recipient_id == user.id:
                continue
            self._notify(
                organization_id,
                recipient_id,
                contract_id,
                "approval_requested",
                "Approval requested",
                f"{request.title} for {contract.title}.",
                f"/contracts/{contract_id}?tab=collaboration",
            )
        self._audit(
            organization_id,
            user.id,
            "approval.requested",
            contract_id,
            {"approval_id": request.id, "assigned_to_user_id": request.assigned_to_user_id},
        )
        self.session.commit()
        self.session.refresh(request)
        return request

    def resolve_approval_request(
        self,
        organization_id: str,
        contract_id: str,
        approval_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> ApprovalRequest:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot decide approval requests.",
        )
        request = self.session.scalar(
            select(ApprovalRequest).where(
                ApprovalRequest.id == approval_id,
                ApprovalRequest.organization_id == organization_id,
                ApprovalRequest.contract_id == contract_id,
            )
        )
        if request is None:
            raise HTTPException(status_code=404, detail="Approval request not found.")
        if request.status != "pending":
            raise HTTPException(status_code=409, detail="This approval request has already been decided.")
        if request.assigned_to_user_id and request.assigned_to_user_id != user.id:
            membership = self.membership(organization_id, user)
            if normalized_role(membership.role) not in {"owner", "admin"}:
                raise HTTPException(status_code=403, detail="This approval is assigned to another reviewer.")
        status_value = payload["status"]
        conditions = json_load(request.conditions_json, [])
        results = payload.get("condition_results", {})
        if status_value in {"approved", "conditionally_approved"} and conditions:
            missing = [condition for condition in conditions if condition not in results]
            if missing:
                raise HTTPException(status_code=422, detail="Record a result for every approval condition.")
            if status_value == "approved" and not all(results.get(condition) for condition in conditions):
                raise HTTPException(status_code=422, detail="Use conditional approval or changes requested while conditions remain unmet.")
        request.status = status_value
        request.condition_results_json = json_dump(results)
        request.resolution_note = payload["resolution_note"].strip()
        request.resolved_by_user_id = user.id
        request.resolved_at = utcnow()
        self._notify(
            organization_id,
            request.requested_by_user_id,
            contract_id,
            "approval_decided",
            f"Approval {status_value.replace('_', ' ')}",
            request.title,
            f"/contracts/{contract_id}?tab=collaboration",
        )
        self._audit(
            organization_id,
            user.id,
            "approval.decided",
            contract_id,
            {"approval_id": request.id, "status": status_value, "condition_results": results},
        )
        self.session.commit()
        self.session.refresh(request)
        return request

    @staticmethod
    def approval_response(request: ApprovalRequest) -> ApprovalRequestResponse:
        return ApprovalRequestResponse(
            id=request.id,
            contract_id=request.contract_id,
            title=request.title,
            note=request.note,
            status=request.status,
            conditions=json_load(request.conditions_json, []),
            condition_results=json_load(request.condition_results_json, {}),
            requested_by_user_id=request.requested_by_user_id,
            requested_by_name=request.requested_by.display_name,
            assigned_to_user_id=request.assigned_to_user_id,
            assigned_to_name=request.assigned_to.display_name if request.assigned_to else None,
            resolved_by_user_id=request.resolved_by_user_id,
            resolved_by_name=request.resolved_by.display_name if request.resolved_by else None,
            resolution_note=request.resolution_note,
            due_at=request.due_at,
            resolved_at=request.resolved_at,
            created_at=request.created_at,
            updated_at=request.updated_at,
        )

    def list_external_shares(self, organization_id: str, contract_id: str, user: User) -> list[ExternalShare]:
        self.get_contract(organization_id, contract_id, user)
        return list(
            self.session.scalars(
                select(ExternalShare)
                .where(
                    ExternalShare.organization_id == organization_id,
                    ExternalShare.contract_id == contract_id,
                )
                .order_by(ExternalShare.created_at.desc())
            ).all()
        )

    def create_external_share(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> tuple[ExternalShare, str]:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can create external links.",
        )
        self.get_review(organization_id, contract_id, user)
        token = secrets.token_urlsafe(32)
        share = ExternalShare(
            organization_id=organization_id,
            contract_id=contract_id,
            created_by_user_id=user.id,
            token_hash=hashlib.sha256(token.encode()).hexdigest(),
            label=payload["label"].strip(),
            include_evidence=payload.get("include_evidence", True),
            expires_at=utcnow() + timedelta(days=payload.get("expires_in_days", 7)),
        )
        self.session.add(share)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "share.created",
            contract_id,
            {"share_id": share.id, "expires_at": share.expires_at, "include_evidence": share.include_evidence},
        )
        self.session.commit()
        self.session.refresh(share)
        return share, token

    def revoke_external_share(
        self,
        organization_id: str,
        contract_id: str,
        share_id: str,
        user: User,
    ) -> None:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin"},
            "Only owners and administrators can revoke external links.",
        )
        share = self.session.scalar(
            select(ExternalShare).where(
                ExternalShare.id == share_id,
                ExternalShare.organization_id == organization_id,
                ExternalShare.contract_id == contract_id,
            )
        )
        if share is None:
            raise HTTPException(status_code=404, detail="Share link not found.")
        share.revoked_at = utcnow()
        self._audit(organization_id, user.id, "share.revoked", contract_id, {"share_id": share.id})
        self.session.commit()

    @staticmethod
    def external_share_response(share: ExternalShare) -> ExternalShareResponse:
        return ExternalShareResponse(
            id=share.id,
            label=share.label,
            include_evidence=share.include_evidence,
            expires_at=share.expires_at,
            revoked_at=share.revoked_at,
            last_viewed_at=share.last_viewed_at,
            view_count=share.view_count,
            created_at=share.created_at,
        )

    def shared_contract(self, token: str) -> tuple[ExternalShare, Contract, dict[str, Any]]:
        share = self.session.scalar(
            select(ExternalShare).where(
                ExternalShare.token_hash == hashlib.sha256(token.encode()).hexdigest()
            )
        )
        if share is None or share.revoked_at is not None or aware(share.expires_at) <= utcnow():
            raise HTTPException(status_code=410, detail="This secure review link is invalid or has expired.")
        contract = self.session.get(Contract, share.contract_id)
        if contract is None or contract.review is None:
            raise HTTPException(status_code=404, detail="Shared review not found.")
        analysis = json_load(contract.review.analysis_json, {})
        if not share.include_evidence:
            risks = []
            for item in analysis.get("risk_assessment", []):
                risks.append({key: value for key, value in item.items() if key not in {"quote", "evidence", "excerpt"}})
            analysis["risk_assessment"] = risks
        share.view_count += 1
        share.last_viewed_at = utcnow()
        self._audit(
            share.organization_id,
            None,
            "share.viewed",
            share.contract_id,
            {"share_id": share.id, "view_count": share.view_count},
        )
        self.session.commit()
        return share, contract, analysis

    def list_lifecycle_items(
        self,
        organization_id: str,
        user: User,
        contract_id: str | None = None,
        status: str | None = None,
    ) -> list[LifecycleItem]:
        self.membership(organization_id, user)
        query = select(LifecycleItem).where(LifecycleItem.organization_id == organization_id)
        if contract_id:
            query = query.where(LifecycleItem.contract_id == contract_id)
        if status:
            query = query.where(LifecycleItem.status == status)
        return list(self.session.scalars(query.order_by(LifecycleItem.due_at.asc())).all())

    def create_lifecycle_item(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        payload: dict[str, Any],
    ) -> LifecycleItem:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot create lifecycle items.",
        )
        self.get_contract(organization_id, contract_id, user)
        owner = self._task_assignee(organization_id, payload.get("owner_user_id"))
        item = LifecycleItem(
            organization_id=organization_id,
            contract_id=contract_id,
            created_by_user_id=user.id,
            owner_user_id=owner.id if owner else None,
            kind=payload["kind"],
            title=payload["title"].strip(),
            description=payload.get("description", "").strip(),
            amount=payload.get("amount", "").strip(),
            due_at=payload["due_at"],
            reminder_days=payload.get("reminder_days", 7),
            recurrence=payload.get("recurrence", "none"),
        )
        self.session.add(item)
        self.session.flush()
        self._audit(
            organization_id,
            user.id,
            "lifecycle.created",
            contract_id,
            {"lifecycle_id": item.id, "kind": item.kind, "due_at": item.due_at, "recurrence": item.recurrence},
        )
        self.session.commit()
        self.session.refresh(item)
        return item

    def update_lifecycle_item(
        self,
        organization_id: str,
        item_id: str,
        user: User,
        changes: dict[str, Any],
    ) -> LifecycleItem:
        self.require_roles(
            organization_id,
            user,
            {"owner", "admin", "reviewer"},
            "Viewers cannot change lifecycle items.",
        )
        item = self.session.scalar(
            select(LifecycleItem).where(
                LifecycleItem.id == item_id,
                LifecycleItem.organization_id == organization_id,
            )
        )
        if item is None:
            raise HTTPException(status_code=404, detail="Lifecycle item not found.")
        if "owner_user_id" in changes:
            owner = self._task_assignee(organization_id, changes["owner_user_id"])
            item.owner_user_id = owner.id if owner else None
        for field in ("title", "description", "amount", "due_at", "reminder_days", "recurrence"):
            if field in changes and changes[field] is not None:
                value = changes[field]
                if isinstance(value, str) and field in {"title", "description", "amount"}:
                    value = value.strip()
                setattr(item, field, value)
        if changes.get("status"):
            item.status = changes["status"]
            item.completed_at = utcnow() if item.status == "completed" else None
            if item.status == "completed" and item.recurrence != "none":
                next_item = LifecycleItem(
                    organization_id=item.organization_id,
                    contract_id=item.contract_id,
                    created_by_user_id=user.id,
                    owner_user_id=item.owner_user_id,
                    kind=item.kind,
                    title=item.title,
                    description=item.description,
                    amount=item.amount,
                    due_at=self._next_occurrence(item.due_at, item.recurrence),
                    reminder_days=item.reminder_days,
                    recurrence=item.recurrence,
                )
                self.session.add(next_item)
        self._audit(
            organization_id,
            user.id,
            "lifecycle.updated",
            item.contract_id,
            {"lifecycle_id": item.id, "changed_fields": sorted(changes)},
        )
        self.session.commit()
        self.session.refresh(item)
        return item

    @staticmethod
    def lifecycle_response(item: LifecycleItem) -> LifecycleItemResponse:
        return LifecycleItemResponse(
            id=item.id,
            organization_id=item.organization_id,
            contract_id=item.contract_id,
            contract_title=item.contract.title,
            kind=item.kind,
            title=item.title,
            description=item.description,
            amount=item.amount,
            due_at=item.due_at,
            owner_user_id=item.owner_user_id,
            owner_name=item.owner.display_name if item.owner else None,
            reminder_days=item.reminder_days,
            recurrence=item.recurrence,
            status=item.status,
            last_notified_at=item.last_notified_at,
            escalated_at=item.escalated_at,
            completed_at=item.completed_at,
            created_at=item.created_at,
            updated_at=item.updated_at,
        )

    @staticmethod
    def _next_occurrence(value: datetime, recurrence: str) -> datetime:
        if recurrence == "weekly":
            return value + timedelta(days=7)
        months = {"monthly": 1, "quarterly": 3, "yearly": 12}.get(recurrence, 0)
        month_index = value.month - 1 + months
        year = value.year + month_index // 12
        month = month_index % 12 + 1
        day = min(value.day, calendar.monthrange(year, month)[1])
        return value.replace(year=year, month=month, day=day)

    def contract_activity(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
        limit: int = 200,
    ) -> list[ContractActivityResponse]:
        self.get_contract(organization_id, contract_id, user)
        rows = self.session.execute(
            select(PlatformAuditEvent, User.display_name, User.email)
            .outerjoin(User, User.id == PlatformAuditEvent.actor_user_id)
            .where(
                PlatformAuditEvent.organization_id == organization_id,
                PlatformAuditEvent.contract_id == contract_id,
            )
            .order_by(PlatformAuditEvent.created_at.desc())
            .limit(min(max(limit, 1), 500))
        ).all()
        return [
            ContractActivityResponse(
                id=event.id,
                action=event.action,
                detail=json_load(event.detail_json, {}),
                actor_user_id=event.actor_user_id,
                actor_name=display_name or email or "Lenslayer system",
                created_at=event.created_at,
            )
            for event, display_name, email in rows
        ]

    def portfolio_question(
        self,
        organization_id: str,
        user: User,
        question: str,
    ) -> PortfolioQuestionResponse:
        self.membership(organization_id, user)
        tokens = {
            token for token in re.findall(r"[a-z0-9]{3,}", question.casefold())
            if token not in {"about", "what", "when", "where", "which", "that", "this", "with", "from", "does", "have"}
        }
        reviews = self.session.execute(
            select(Contract, ContractReview)
            .join(ContractReview, ContractReview.contract_id == Contract.id)
            .where(Contract.organization_id == organization_id)
        ).all()
        candidates: list[tuple[int, Contract, str, str]] = []
        for contract, review in reviews:
            analysis = json_load(review.analysis_json, {})
            text = review.source_text or json_dump(analysis)
            blocks = [item.strip() for item in re.split(r"\n{2,}", text) if item.strip()]
            for index, block in enumerate(blocks[:80]):
                score = sum(block.casefold().count(token) for token in tokens)
                if score:
                    candidates.append((score, contract, f"Excerpt {index + 1}", " ".join(block.split())[:900]))
        selected = sorted(candidates, key=lambda item: item[0], reverse=True)[:8]
        sources = [
            PortfolioQuestionSource(
                contract_id=contract.id,
                contract_title=contract.title,
                location=location,
                excerpt=excerpt,
            )
            for _, contract, location, excerpt in selected
        ]
        if not sources:
            return PortfolioQuestionResponse(
                answer="No retained contract evidence matched that question.",
                sources=[],
                generated_by="extractive",
            )
        evidence = "\n\n".join(
            f"[Source {index}] {source.contract_title}, {source.location}\n{source.excerpt}"
            for index, source in enumerate(sources, 1)
        )
        if os.environ.get("GROQ_API_KEY"):
            from analyzer import get_llm

            response = get_llm().invoke(
                "Answer the portfolio question using only the evidence below. Cite [Source 1] inline, "
                "name the relevant contract, and say when evidence is insufficient. This is operational "
                "contract triage, not legal advice.\n\n"
                f"QUESTION\n{question}\n\nEVIDENCE\n{evidence}"
            )
            answer, generated_by = str(response.content).strip(), "model"
        else:
            answer, generated_by = (
                f"Found relevant evidence across {len({source.contract_id for source in sources})} contract(s). "
                "Review the cited excerpts below before acting.",
                "extractive",
            )
        self._audit(
            organization_id,
            user.id,
            "portfolio.question_answered",
            detail={"question": question[:240], "source_count": len(sources), "generated_by": generated_by},
        )
        self.session.commit()
        return PortfolioQuestionResponse(answer=answer, sources=sources, generated_by=generated_by)

    def counsel_handoff(
        self,
        organization_id: str,
        contract_id: str,
        user: User,
    ) -> tuple[bytes, str]:
        contract = self.get_contract(organization_id, contract_id, user)
        review = self.get_review(organization_id, contract_id, user)
        from docx import Document

        analysis = json_load(review.analysis_json, {})
        document = Document()
        document.add_heading(f"Counsel handoff: {contract.title}", 0)
        document.add_paragraph("Prepared by Lenslayer for qualified professional review. Not legal advice.")
        document.add_heading("Executive summary", 1)
        document.add_paragraph(str(analysis.get("executive_summary") or "No executive summary returned."))
        document.add_heading("Open risks and protection gaps", 1)
        for item in analysis.get("risk_assessment", []):
            document.add_heading(str(item.get("title") or "Clause finding"), 2)
            document.add_paragraph(str(item.get("explanation") or ""))
            document.add_paragraph(f"Evidence: {item.get('citation') or 'Not identified'} | {item.get('quote') or 'No quote returned'}")
        for item in analysis.get("missing_protections", []):
            document.add_paragraph(str(item.get("issue") if isinstance(item, dict) else item), style="List Bullet")
        document.add_heading("Negotiation priorities", 1)
        for item in analysis.get("negotiation_priorities", []):
            document.add_paragraph(str(item.get("title") if isinstance(item, dict) else item), style="List Bullet")
        document.add_heading("Human decisions", 1)
        for decision in self.list_contract_decisions(organization_id, contract_id, user):
            document.add_paragraph(
                f"{decision.decision.upper()}: {decision.subject} | {decision.reviewer.display_name} | {decision.rationale}",
                style="List Bullet",
            )
        document.add_heading("Open actions and approvals", 1)
        for task in self.list_tasks(organization_id, user, contract_id=contract_id):
            if task.status not in {"done", "cancelled"}:
                document.add_paragraph(f"{task.title} ({task.status})", style="List Bullet")
        for approval in self.list_approval_requests(organization_id, contract_id, user):
            document.add_paragraph(f"{approval.title} ({approval.status})", style="List Bullet")
        output = io.BytesIO()
        document.save(output)
        self._audit(organization_id, user.id, "counsel_handoff.exported", contract_id)
        self.session.commit()
        return output.getvalue(), f"{safe_filename(contract.title)}-counsel-handoff.docx"

    def calendar_ics(self, organization_id: str, user: User) -> str:
        self.membership(organization_id, user)
        events: list[tuple[str, str, datetime, str]] = []
        for task in self.list_tasks(organization_id, user):
            if task.due_at and task.status not in {"done", "cancelled"}:
                events.append((f"task-{task.id}", task.title, task.due_at, task.contract.title if task.contract else "Lenslayer"))
        for item in self.list_lifecycle_items(organization_id, user, status="active"):
            events.append((f"lifecycle-{item.id}", item.title, item.due_at, f"{item.kind}: {item.contract.title}"))
        def escape(value: str) -> str:
            return value.replace("\\", "\\\\").replace(",", "\\,").replace(";", "\\;").replace("\n", "\\n")
        lines = ["BEGIN:VCALENDAR", "VERSION:2.0", "PRODID:-//Lenslayer//Lifecycle Calendar//EN", "CALSCALE:GREGORIAN"]
        for uid, title, due_at, description in events:
            stamp = aware(due_at).astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
            lines.extend([
                "BEGIN:VEVENT",
                f"UID:{uid}@lenslayer",
                f"DTSTAMP:{utcnow().strftime('%Y%m%dT%H%M%SZ')}",
                f"DTSTART:{stamp}",
                f"SUMMARY:{escape(title)}",
                f"DESCRIPTION:{escape(description)}",
                "END:VEVENT",
            ])
        lines.append("END:VCALENDAR")
        return "\r\n".join(lines) + "\r\n"

    def _notify(
        self,
        organization_id: str,
        user_id: str,
        contract_id: str | None,
        kind: str,
        title: str,
        message: str,
        action_url: str,
    ) -> None:
        self.session.add(
            Notification(
                organization_id=organization_id,
                user_id=user_id,
                contract_id=contract_id,
                kind=kind,
                title=title,
                message=message,
                action_url=action_url,
            )
        )

    def list_notifications(
        self,
        organization_id: str,
        user: User,
        unread_only: bool = False,
        limit: int = 30,
    ) -> list[Notification]:
        self.membership(organization_id, user)
        query = select(Notification).where(
            Notification.organization_id == organization_id,
            Notification.user_id == user.id,
        )
        if unread_only:
            query = query.where(Notification.read_at.is_(None))
        return list(
            self.session.scalars(
                query.order_by(Notification.created_at.desc()).limit(min(max(limit, 1), 100))
            ).all()
        )

    def mark_notification_read(
        self,
        organization_id: str,
        notification_id: str,
        user: User,
    ) -> Notification:
        self.membership(organization_id, user)
        notification = self.session.scalar(
            select(Notification).where(
                Notification.id == notification_id,
                Notification.organization_id == organization_id,
                Notification.user_id == user.id,
            )
        )
        if notification is None:
            raise HTTPException(status_code=404, detail="Notification not found.")
        if notification.read_at is None:
            notification.read_at = utcnow()
            self.session.commit()
            self.session.refresh(notification)
        return notification

    def mark_all_notifications_read(self, organization_id: str, user: User) -> None:
        for notification in self.list_notifications(organization_id, user, unread_only=True, limit=100):
            notification.read_at = utcnow()
        self.session.commit()

    @staticmethod
    def notification_response(notification: Notification) -> NotificationResponse:
        return NotificationResponse(
            id=notification.id,
            organization_id=notification.organization_id,
            contract_id=notification.contract_id,
            kind=notification.kind,
            title=notification.title,
            message=notification.message,
            action_url=notification.action_url,
            read_at=notification.read_at,
            created_at=notification.created_at,
        )

    def list_jobs(self, organization_id: str, contract_id: str, user: User) -> list[ProcessingJob]:
        self.get_contract(organization_id, contract_id, user)
        return list(
            self.session.scalars(
                select(ProcessingJob)
                .where(
                    ProcessingJob.organization_id == organization_id,
                    ProcessingJob.contract_id == contract_id,
                )
                .order_by(ProcessingJob.created_at.desc())
            ).all()
        )

    def list_audit_events(self, organization_id: str, user: User, limit: int = 100) -> list[PlatformAuditEvent]:
        self.membership(organization_id, user)
        return list(
            self.session.scalars(
                select(PlatformAuditEvent)
                .where(PlatformAuditEvent.organization_id == organization_id)
                .order_by(PlatformAuditEvent.created_at.desc())
                .limit(min(max(limit, 1), 250))
            ).all()
        )

    def report_overview(
        self,
        organization_id: str,
        user: User,
        range_name: str,
    ) -> ReportOverviewResponse:
        self.membership(organization_id, user)
        if range_name not in {"30d", "90d", "365d", "all"}:
            raise HTTPException(status_code=422, detail="Choose a 30 day, 90 day, 365 day, or all-time report.")

        generated_at = utcnow()
        range_days = {"30d": 30, "90d": 90, "365d": 365}
        period_start = generated_at - timedelta(days=range_days[range_name]) if range_name != "all" else None

        contracts = list(
            self.session.scalars(
                select(Contract)
                .where(Contract.organization_id == organization_id)
                .order_by(Contract.created_at.asc())
            ).all()
        )
        tasks = list(
            self.session.scalars(
                select(WorkflowTask)
                .where(WorkflowTask.organization_id == organization_id)
                .order_by(WorkflowTask.created_at.asc())
            ).all()
        )
        verification_cases = list(
            self.session.scalars(
                select(VerificationCase)
                .where(VerificationCase.organization_id == organization_id)
                .order_by(VerificationCase.submitted_at.asc())
            ).all()
        )
        decisions = list(
            self.session.scalars(
                select(VerificationDecision)
                .where(VerificationDecision.organization_id == organization_id)
                .order_by(VerificationDecision.created_at.asc())
            ).all()
        )
        audit_events = list(
            self.session.scalars(
                select(PlatformAuditEvent)
                .where(PlatformAuditEvent.organization_id == organization_id)
                .order_by(PlatformAuditEvent.created_at.desc())
            ).all()
        )
        memberships = list(
            self.session.scalars(
                select(Membership)
                .join(User, User.id == Membership.user_id)
                .where(Membership.organization_id == organization_id)
                .order_by(User.display_name.asc(), User.email.asc())
            ).all()
        )

        def in_period(value: datetime) -> bool:
            return period_start is None or aware(value) >= period_start

        period_contracts = [item for item in contracts if in_period(item.created_at)]
        period_tasks = [item for item in tasks if in_period(item.created_at)]
        period_cases = [item for item in verification_cases if in_period(item.submitted_at)]
        period_decisions = [item for item in decisions if in_period(item.created_at)]
        period_events = [item for item in audit_events if in_period(item.created_at)]

        active_tasks = [item for item in tasks if item.status in {"open", "in_progress"}]
        overdue_tasks = [
            item for item in active_tasks
            if item.due_at is not None and aware(item.due_at) < generated_at
        ]
        due_soon_cutoff = generated_at + timedelta(days=7)
        due_soon_tasks = [
            item for item in active_tasks
            if item.due_at is not None and generated_at <= aware(item.due_at) <= due_soon_cutoff
        ]
        completed_in_period = [
            item for item in tasks
            if item.status == "done" and item.completed_at is not None and in_period(item.completed_at)
        ]
        period_completed_tasks = [item for item in period_tasks if item.status == "done"]
        task_denominator = sum(1 for item in period_tasks if item.status != "cancelled")

        contract_type_counts: dict[str, int] = {}
        for contract in period_contracts:
            label = contract.contract_type.strip() or "Unknown"
            contract_type_counts[label] = contract_type_counts.get(label, 0) + 1

        priority_counts = {
            priority: sum(1 for task in active_tasks if task.priority == priority)
            for priority in ("high", "normal", "low")
        }

        timeline_start = period_start or self._earliest_report_time(
            generated_at,
            contracts,
            tasks,
            verification_cases,
            decisions,
        )
        timeline = self._report_timeline(
            timeline_start,
            generated_at,
            contracts,
            tasks,
            verification_cases,
            decisions,
            range_name,
        )

        workload = [
            ReportWorkloadItem(
                user_id=membership.user_id,
                display_name=membership.user.display_name,
                email=membership.user.email,
                role=normalized_role(membership.role),
                active_tasks=sum(
                    1 for task in active_tasks if task.assigned_to_user_id == membership.user_id
                ),
                overdue_tasks=sum(
                    1 for task in overdue_tasks if task.assigned_to_user_id == membership.user_id
                ),
                completed_in_period=sum(
                    1 for task in completed_in_period if task.assigned_to_user_id == membership.user_id
                ),
            )
            for membership in memberships
        ]

        actor_ids = {event.actor_user_id for event in period_events[:12] if event.actor_user_id}
        actors = {
            actor.id: actor
            for actor in self.session.scalars(select(User).where(User.id.in_(actor_ids))).all()
        } if actor_ids else {}
        contract_ids = {event.contract_id for event in period_events[:12] if event.contract_id}
        activity_contracts = {
            contract.id: contract
            for contract in self.session.scalars(select(Contract).where(Contract.id.in_(contract_ids))).all()
        } if contract_ids else {}
        recent_activity = [
            ReportActivityItem(
                id=event.id,
                action=event.action,
                detail=json_load(event.detail_json, {}),
                actor_user_id=event.actor_user_id,
                actor_name=actors[event.actor_user_id].display_name
                if event.actor_user_id in actors
                else "Lenslayer system",
                contract_id=event.contract_id,
                contract_title=activity_contracts[event.contract_id].title
                if event.contract_id in activity_contracts
                else None,
                created_at=event.created_at,
            )
            for event in period_events[:12]
        ]

        return ReportOverviewResponse(
            organization_id=organization_id,
            range=range_name,
            generated_at=generated_at,
            period_start=period_start,
            period_end=generated_at,
            contracts_total=len(period_contracts),
            contracts_ready=sum(1 for item in period_contracts if item.status == "ready"),
            contracts_processing=sum(
                1 for item in period_contracts if item.status in {"queued", "processing", "running"}
            ),
            contracts_failed=sum(1 for item in period_contracts if item.status == "failed"),
            tasks_total=len(period_tasks),
            tasks_active=len(active_tasks),
            tasks_overdue=len(overdue_tasks),
            tasks_due_soon=len(due_soon_tasks),
            tasks_completed=len(completed_in_period),
            task_completion_rate=round((len(period_completed_tasks) / task_denominator) * 100)
            if task_denominator
            else 0,
            verification_total=len(period_cases),
            verification_pending=sum(1 for item in period_cases if item.status == "pending"),
            verification_approved=sum(1 for item in period_cases if item.status == "approved"),
            verification_escalated=sum(1 for item in period_cases if item.status == "escalated"),
            verification_rejected=sum(1 for item in period_cases if item.status == "rejected"),
            verification_average_risk=round(
                sum(item.risk_score for item in period_cases) / len(period_cases)
            ) if period_cases else 0,
            verification_overrides=sum(
                1 for item in period_decisions if item.decision != item.recommended_action
            ),
            audit_event_count=len(period_events),
            contract_types=[
                ReportDistributionItem(label=label, count=count)
                for label, count in sorted(
                    contract_type_counts.items(),
                    key=lambda item: (-item[1], item[0].casefold()),
                )
            ],
            active_task_priorities=[
                ReportDistributionItem(label=priority, count=priority_counts[priority])
                for priority in ("high", "normal", "low")
            ],
            timeline=timeline,
            workload=sorted(
                workload,
                key=lambda item: (-item.overdue_tasks, -item.active_tasks, item.display_name.casefold()),
            ),
            recent_activity=recent_activity,
        )

    @staticmethod
    def report_csv(report: ReportOverviewResponse) -> str:
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["Lenslayer report", report.range, report.generated_at.isoformat()])
        writer.writerow([])
        writer.writerow(["Section", "Metric", "Value"])
        metrics = [
            ("Contracts", "Created", report.contracts_total),
            ("Contracts", "Ready", report.contracts_ready),
            ("Contracts", "Processing", report.contracts_processing),
            ("Contracts", "Failed", report.contracts_failed),
            ("Tasks", "Created", report.tasks_total),
            ("Tasks", "Currently active", report.tasks_active),
            ("Tasks", "Currently overdue", report.tasks_overdue),
            ("Tasks", "Due in seven days", report.tasks_due_soon),
            ("Tasks", "Completed in period", report.tasks_completed),
            ("Tasks", "Completion rate", f"{report.task_completion_rate}%"),
            ("Verify", "Cases submitted", report.verification_total),
            ("Verify", "Pending", report.verification_pending),
            ("Verify", "Approved", report.verification_approved),
            ("Verify", "Escalated", report.verification_escalated),
            ("Verify", "Rejected", report.verification_rejected),
            ("Verify", "Average risk score", report.verification_average_risk),
            ("Verify", "Recommendation overrides", report.verification_overrides),
            ("Governance", "Audit events", report.audit_event_count),
        ]
        writer.writerows(metrics)
        writer.writerow([])
        writer.writerow(["Reviewer", "Email", "Role", "Active tasks", "Overdue tasks", "Completed in period"])
        for item in report.workload:
            writer.writerow([
                item.display_name,
                item.email,
                item.role,
                item.active_tasks,
                item.overdue_tasks,
                item.completed_in_period,
            ])
        writer.writerow([])
        writer.writerow([
            "Period",
            "Contracts created",
            "Tasks created",
            "Tasks completed",
            "Verification cases",
            "Decisions",
        ])
        for point in report.timeline:
            writer.writerow([
                point.label,
                point.contracts_created,
                point.tasks_created,
                point.tasks_completed,
                point.verification_submitted,
                point.decisions_recorded,
            ])
        return output.getvalue()

    @staticmethod
    def _earliest_report_time(
        fallback: datetime,
        contracts: list[Contract],
        tasks: list[WorkflowTask],
        verification_cases: list[VerificationCase],
        decisions: list[VerificationDecision],
    ) -> datetime:
        values = (
            [aware(item.created_at) for item in contracts]
            + [aware(item.created_at) for item in tasks]
            + [aware(item.submitted_at) for item in verification_cases]
            + [aware(item.created_at) for item in decisions]
        )
        return min(values) if values else fallback - timedelta(days=30)

    @staticmethod
    def _report_timeline(
        period_start: datetime,
        period_end: datetime,
        contracts: list[Contract],
        tasks: list[WorkflowTask],
        verification_cases: list[VerificationCase],
        decisions: list[VerificationDecision],
        range_name: str,
    ) -> list[ReportTimelinePoint]:
        configured_days = {"30d": 5, "90d": 15, "365d": 31}
        if range_name == "all":
            span_days = max(1, math.ceil((period_end - period_start).total_seconds() / 86400))
            bucket_days = max(1, math.ceil(span_days / 8))
        else:
            bucket_days = configured_days[range_name]
        points: list[ReportTimelinePoint] = []
        cursor = period_start
        while cursor < period_end:
            bucket_end = min(cursor + timedelta(days=bucket_days), period_end)

            def inside(value: datetime | None) -> bool:
                return value is not None and cursor <= aware(value) < bucket_end

            points.append(
                ReportTimelinePoint(
                    label=cursor.strftime("%d %b"),
                    period_start=cursor,
                    period_end=bucket_end,
                    contracts_created=sum(1 for item in contracts if inside(item.created_at)),
                    tasks_created=sum(1 for item in tasks if inside(item.created_at)),
                    tasks_completed=sum(1 for item in tasks if inside(item.completed_at)),
                    verification_submitted=sum(
                        1 for item in verification_cases if inside(item.submitted_at)
                    ),
                    decisions_recorded=sum(1 for item in decisions if inside(item.created_at)),
                )
            )
            cursor = bucket_end
        return points

    def delete_contract(self, organization_id: str, contract_id: str, user: User) -> None:
        membership = self.membership(organization_id, user)
        if membership.role not in {"owner", "admin"}:
            raise HTTPException(status_code=403, detail="Only owners and administrators can delete contracts.")
        contract = self.get_contract(organization_id, contract_id, user)
        keys = [asset.storage_key for asset in contract.assets]
        self._audit(
            organization_id,
            user.id,
            "contract.deleted",
            None,
            {"contract_id": contract.id, "title": contract.title},
        )
        self.session.delete(contract)
        self.session.commit()
        for key in keys:
            self.object_store.delete(key)

    def contract_response(self, contract: Contract) -> ContractResponse:
        latest = max(contract.jobs, key=lambda item: item.created_at) if contract.jobs else None
        return ContractResponse(
            id=contract.id,
            organization_id=contract.organization_id,
            title=contract.title,
            source_name=contract.source_name,
            counterparty=contract.counterparty,
            contract_type=contract.contract_type,
            status=contract.status,
            review_context=json_load(contract.review_context_json, {}),
            retain_document=contract.retain_document,
            retain_source_text=contract.retain_source_text,
            retention_days=contract.retention_days,
            expires_at=contract.expires_at,
            created_at=contract.created_at,
            updated_at=contract.updated_at,
            latest_job=JobResponse.model_validate(latest) if latest else None,
        )

    @staticmethod
    def integration_connection_response(connection: IntegrationConnection) -> IntegrationConnectionResponse:
        return IntegrationConnectionResponse(
            id=connection.id,
            organization_id=connection.organization_id,
            provider=connection.provider,
            display_name=connection.display_name,
            external_account_id=connection.external_account_id,
            status=connection.status,
            capabilities=json_load(connection.capabilities_json, []),
            settings=json_load(connection.settings_json, {}),
            last_sync_at=connection.last_sync_at,
            error_message=connection.error_message,
            created_by_user_id=connection.created_by_user_id,
            created_by_name=connection.created_by.display_name,
            created_at=connection.created_at,
            updated_at=connection.updated_at,
        )

    @staticmethod
    def integration_import_response(import_record: IntegrationImport) -> IntegrationImportResponse:
        return IntegrationImportResponse(
            id=import_record.id,
            organization_id=import_record.organization_id,
            connection_id=import_record.connection_id,
            contract_id=import_record.contract_id,
            provider=import_record.provider,
            source_type=import_record.source_type,
            external_id=import_record.external_id,
            source_url=import_record.source_url,
            title=import_record.title,
            original_name=import_record.original_name,
            content_type=import_record.content_type,
            size_bytes=import_record.size_bytes,
            sha256=import_record.sha256,
            status=import_record.status,
            metadata=json_load(import_record.metadata_json, {}),
            error_message=import_record.error_message,
            imported_by_user_id=import_record.imported_by_user_id,
            imported_by_name=import_record.imported_by.display_name if import_record.imported_by else "API key",
            created_at=import_record.created_at,
            updated_at=import_record.updated_at,
        )

    @staticmethod
    def api_key_response(api_key: PublicApiKey) -> ApiKeyResponse:
        return ApiKeyResponse(
            id=api_key.id,
            organization_id=api_key.organization_id,
            name=api_key.name,
            key_prefix=api_key.key_prefix,
            scopes=json_load(api_key.scopes_json, []),
            last_used_at=api_key.last_used_at,
            revoked_at=api_key.revoked_at,
            created_by_user_id=api_key.created_by_user_id,
            created_by_name=api_key.created_by.display_name,
            created_at=api_key.created_at,
        )

    @staticmethod
    def webhook_subscription_response(webhook: WebhookSubscription) -> WebhookSubscriptionResponse:
        return WebhookSubscriptionResponse(
            id=webhook.id,
            organization_id=webhook.organization_id,
            target_url=webhook.target_url,
            description=webhook.description,
            events=json_load(webhook.events_json, []),
            secret_prefix=webhook.secret_prefix,
            status=webhook.status,
            last_delivery_at=webhook.last_delivery_at,
            created_by_user_id=webhook.created_by_user_id,
            created_by_name=webhook.created_by.display_name,
            created_at=webhook.created_at,
            updated_at=webhook.updated_at,
        )

    @staticmethod
    def webhook_delivery_response(delivery: WebhookDelivery) -> WebhookDeliveryResponse:
        return WebhookDeliveryResponse(
            id=delivery.id,
            subscription_id=delivery.subscription_id,
            contract_id=delivery.contract_id,
            event_type=delivery.event_type,
            payload=json_load(delivery.payload_json, {}),
            status=delivery.status,
            attempts=delivery.attempts,
            last_error=delivery.last_error,
            delivered_at=delivery.delivered_at,
            created_at=delivery.created_at,
            updated_at=delivery.updated_at,
        )

    @staticmethod
    def membership_response(membership: Membership) -> MembershipResponse:
        return MembershipResponse(
            id=membership.id,
            user_id=membership.user_id,
            email=membership.user.email,
            display_name=membership.user.display_name,
            role=normalized_role(membership.role),
            created_at=membership.created_at,
        )

    @staticmethod
    def invitation_response(invitation: OrganizationInvitation) -> InvitationResponse:
        return InvitationResponse(
            id=invitation.id,
            email=invitation.email,
            role=invitation.role,
            status=invitation_status(invitation),
            expires_at=invitation.expires_at,
            accepted_at=invitation.accepted_at,
            created_at=invitation.created_at,
        )

    @staticmethod
    def invitation_preview(invitation: OrganizationInvitation) -> InvitationPreviewResponse:
        return InvitationPreviewResponse(
            organization_name=invitation.organization.name,
            organization_slug=invitation.organization.slug,
            email_hint=email_hint(invitation.email),
            role=invitation.role,
            status=invitation_status(invitation),
            expires_at=invitation.expires_at,
        )

    @staticmethod
    def task_response(task: WorkflowTask) -> TaskResponse:
        return TaskResponse(
            id=task.id,
            organization_id=task.organization_id,
            contract_id=task.contract_id,
            contract_title=task.contract.title if task.contract else None,
            created_by_user_id=task.created_by_user_id,
            assigned_to_user_id=task.assigned_to_user_id,
            assigned_to_name=task.assigned_to_user.display_name if task.assigned_to_user else None,
            assigned_to_email=task.assigned_to_user.email if task.assigned_to_user else None,
            title=task.title,
            description=task.description,
            category=task.category,
            priority=task.priority,
            status=task.status,
            due_at=task.due_at,
            source_kind=task.source_kind,
            source_reference=json_load(task.source_reference_json, {}),
            completed_at=task.completed_at,
            created_at=task.created_at,
            updated_at=task.updated_at,
        )

    @staticmethod
    def verification_decision_response(decision: VerificationDecision) -> VerificationDecisionResponse:
        return VerificationDecisionResponse(
            id=decision.id,
            decision=decision.decision,
            rationale=decision.rationale,
            recommended_action=decision.recommended_action,
            reviewer_user_id=decision.reviewer_user_id,
            reviewer_name=decision.reviewer.display_name,
            reviewer_email=decision.reviewer.email,
            created_at=decision.created_at,
        )

    @staticmethod
    def verification_document_response(document: VerificationDocument) -> VerificationDocumentResponse:
        return VerificationDocumentResponse(
            id=document.id,
            document_type=document.document_type,
            original_name=document.original_name,
            content_type=document.content_type,
            size_bytes=document.size_bytes,
            sha256=document.sha256,
            status=document.status,
            scan_status=document.scan_status,
            extraction_status=document.extraction_status,
            extracted_fields=json_load(document.extracted_fields_json, {}),
            confidence=document.confidence,
            uploaded_by_user_id=document.uploaded_by_user_id,
            uploaded_by_name=document.uploaded_by.display_name if document.uploaded_by else "Secure intake",
            expires_at=document.expires_at,
            created_at=document.created_at,
            updated_at=document.updated_at,
        )

    @staticmethod
    def verification_assignment_response(assignment: VerificationAssignment) -> VerificationAssignmentResponse:
        return VerificationAssignmentResponse(
            id=assignment.id,
            assigned_to_user_id=assignment.assigned_to_user_id,
            assigned_to_name=assignment.assigned_to.display_name if assignment.assigned_to else "Unassigned",
            assigned_to_email=assignment.assigned_to.email if assignment.assigned_to else "",
            assigned_by_user_id=assignment.assigned_by_user_id,
            assigned_by_name=assignment.assigned_by.display_name,
            note=assignment.note,
            created_at=assignment.created_at,
        )

    @staticmethod
    def verification_reconciliation_response(
        record: VerificationReconciliation,
    ) -> VerificationReconciliationResponse:
        return VerificationReconciliationResponse(
            id=record.id,
            field_name=record.field_name,
            canonical_value=record.canonical_value,
            status=record.status,
            sources=json_load(record.sources_json, []),
            resolution_note=record.resolution_note,
            resolved_by_user_id=record.resolved_by_user_id,
            resolved_by_name=record.resolved_by.display_name if record.resolved_by else "",
            resolved_at=record.resolved_at,
            created_at=record.created_at,
            updated_at=record.updated_at,
        )

    def verification_case_summary_response(
        self,
        case: VerificationCase,
    ) -> VerificationCaseSummaryResponse:
        latest = case.decisions[-1] if case.decisions else None
        return VerificationCaseSummaryResponse(
            id=case.id,
            organization_id=case.organization_id,
            reference=case.reference,
            applicant_name=case.applicant_name,
            applicant_email=case.applicant_email,
            status=case.status,
            priority=case.priority,
            assigned_to_user_id=case.assigned_to_user_id,
            assigned_to_name=case.assigned_to.display_name if case.assigned_to else "",
            assigned_to_email=case.assigned_to.email if case.assigned_to else "",
            intake_channel=case.intake_channel,
            risk_score=case.risk_score,
            suggested_action=case.suggested_action,
            finding_count=case.finding_count,
            document_count=case.document_count,
            average_confidence=case.average_confidence,
            submitted_at=case.submitted_at,
            synthetic=case.synthetic,
            due_at=case.due_at,
            expires_at=case.expires_at,
            closed_at=case.closed_at,
            latest_decision=self.verification_decision_response(latest) if latest else None,
            created_at=case.created_at,
            updated_at=case.updated_at,
        )

    def verification_case_response(self, case: VerificationCase) -> VerificationCaseResponse:
        source = json_load(case.source_json, {})
        evaluation = json_load(case.evaluation_json, {})
        summary = self.verification_case_summary_response(case)
        return VerificationCaseResponse(
            **summary.model_dump(),
            application=source.get("application", {}),
            documents=source.get("documents", []),
            summary=evaluation.get("summary", ""),
            reasoning=evaluation.get("reasoning", ""),
            findings=evaluation.get("findings", []),
            field_matrix=evaluation.get("field_matrix", []),
            generated_at=evaluation.get("generated_at", ""),
            decision_history=[self.verification_decision_response(item) for item in case.decisions],
            uploaded_documents=[self.verification_document_response(item) for item in case.documents],
            assignment_history=[self.verification_assignment_response(item) for item in case.assignments],
            reconciliations=[
                self.verification_reconciliation_response(item) for item in case.reconciliations
            ],
        )

    def secure_intake_link_response(self, link: SecureIntakeLink) -> SecureIntakeLinkResponse:
        return SecureIntakeLinkResponse(
            id=link.id,
            organization_id=link.organization_id,
            verification_case_id=link.verification_case_id,
            token_prefix=link.token_prefix,
            channel=link.channel,
            recipient_name=link.recipient_name,
            recipient_email=link.recipient_email,
            recipient_phone_hint=link.recipient_phone_hint,
            applicant_name=link.applicant_name,
            message=link.message,
            max_uploads=link.max_uploads,
            upload_count=link.upload_count,
            retention_days=link.retention_days,
            status=self.secure_intake_status(link),
            expires_at=link.expires_at,
            revoked_at=link.revoked_at,
            last_used_at=link.last_used_at,
            created_by_user_id=link.created_by_user_id,
            created_by_name=link.created_by.display_name,
            created_at=link.created_at,
        )

    @staticmethod
    def review_response(review: ContractReview) -> ReviewResponse:
        return ReviewResponse(
            id=review.id,
            analysis=json_load(review.analysis_json, {}),
            quality=json_load(review.quality_json, {}),
            source_text_retained=bool(review.source_text),
            created_at=review.created_at,
            updated_at=review.updated_at,
        )

    def _audit(
        self,
        organization_id: str,
        actor_user_id: str | None,
        action: str,
        contract_id: str | None = None,
        detail: dict[str, Any] | None = None,
        verification_case_id: str | None = None,
    ) -> None:
        self.session.add(
            PlatformAuditEvent(
                organization_id=organization_id,
                actor_user_id=actor_user_id,
                contract_id=contract_id,
                verification_case_id=verification_case_id,
                action=action,
                detail_json=json_dump(detail or {}),
            )
        )
