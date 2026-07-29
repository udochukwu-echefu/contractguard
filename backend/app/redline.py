from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _candidate(finding: dict[str, Any]) -> str:
    for key in ("clause", "evidence", "excerpt", "quote"):
        value = str(finding.get(key) or "").strip()
        if value:
            return " ".join(value.split())
    return ""


def _tracked_text(paragraph: Any, old_text: str, new_text: str, change_id: int) -> None:
    paragraph_element = paragraph._p
    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)
    timestamp = datetime.now(timezone.utc).isoformat()
    if old_text:
        deleted = OxmlElement("w:del")
        deleted.set(qn("w:id"), str(change_id))
        deleted.set(qn("w:author"), "Lenslayer")
        deleted.set(qn("w:date"), timestamp)
        run = OxmlElement("w:r")
        text = OxmlElement("w:delText")
        text.set(qn("xml:space"), "preserve")
        text.text = old_text
        run.append(text)
        deleted.append(run)
        paragraph_element.append(deleted)
    inserted = OxmlElement("w:ins")
    inserted.set(qn("w:id"), str(change_id + 1))
    inserted.set(qn("w:author"), "Lenslayer")
    inserted.set(qn("w:date"), timestamp)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.set(qn("xml:space"), "preserve")
    text.text = new_text
    run.append(text)
    inserted.append(run)
    paragraph_element.append(inserted)


def build_redline(source: bytes, findings: list[dict[str, Any]]) -> tuple[bytes, int]:
    document = Document(io.BytesIO(source))
    changes = [
        finding
        for finding in findings
        if isinstance(finding, dict) and str(finding.get("suggested_language") or "").strip()
    ]
    matched_ids: set[int] = set()
    change_id = 1
    for index, finding in enumerate(changes):
        needle = _candidate(finding)
        if len(needle) < 12:
            continue
        for paragraph in document.paragraphs:
            current = " ".join(paragraph.text.split())
            if needle.casefold() in current.casefold() or current.casefold() in needle.casefold():
                original = paragraph.text
                replacement = str(finding["suggested_language"]).strip()
                _tracked_text(paragraph, original, replacement, change_id)
                note = paragraph.add_run(" ")
                document.add_comment(
                    note,
                    text=str(finding.get("recommendation") or finding.get("explanation") or "Review this proposed wording."),
                    author="Lenslayer",
                    initials="LL",
                )
                matched_ids.add(index)
                change_id += 2
                break
    unmatched = [(index, finding) for index, finding in enumerate(changes) if index not in matched_ids]
    if unmatched:
        document.add_heading("Lenslayer proposed changes", level=1)
        document.add_paragraph(
            "These suggestions could not be anchored to an exact paragraph. Review and place them before accepting."
        )
        for _, finding in unmatched:
            paragraph = document.add_paragraph()
            _tracked_text(paragraph, "", str(finding["suggested_language"]).strip(), change_id)
            note = paragraph.add_run(" ")
            document.add_comment(
                note,
                text=str(finding.get("recommendation") or finding.get("title") or "Review this proposed wording."),
                author="Lenslayer",
                initials="LL",
            )
            change_id += 2
    output = io.BytesIO()
    document.save(output)
    return output.getvalue(), len(changes)
